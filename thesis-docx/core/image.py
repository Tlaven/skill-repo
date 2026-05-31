"""
core/image.py — Image write operations.

Mixin for SafeDocument. Methods assume self.model, self._rebuild(),
self._media_overrides, self._detect_caption_style() exist.
"""

from __future__ import annotations
import os

from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .types import Anchor


class ImageMixin:

    def insert_image(self, after: Anchor, image_path: str,
                     width: float | None = None, caption: str | None = None) -> bool:
        """Insert an image after the specified paragraph. Does NOT call save()."""
        if after.kind != "paragraph" or after.paragraph_index is None:
            return False
        if not os.path.isfile(image_path):
            return False

        ext = os.path.splitext(image_path)[1].lower()
        if ext == ".svg":
            return False

        idx = after.paragraph_index
        if idx < 0 or idx >= len(self.model._doc.paragraphs):
            return False

        section = self.model._doc.sections[0]
        text_width_emu = section.page_width - section.left_margin - section.right_margin
        text_width_cm = text_width_emu / 914400 * 2.54
        max_height_cm = (section.page_height - section.top_margin - section.bottom_margin) / 914400 * 2.54 - 2

        width_cm = min(width or text_width_cm * 0.8, text_width_cm)

        p = self.model._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        inline = run.add_picture(image_path, width=Cm(width_cm))

        # Maintain aspect ratio if too tall
        if inline.height and max_height_cm > 0:
            height_cm = inline.height / 914400 * 2.54
            if height_cm > max_height_cm:
                scale = max_height_cm / height_cm
                inline.width = int(inline.width * scale)
                inline.height = int(inline.height * scale)

        # Reposition from default location
        p._element.getparent().remove(p._element)
        ref_elem = self.model._doc.paragraphs[idx]._element
        ref_elem.addnext(p._element)

        if caption:
            caption_style = self._detect_caption_style()
            cap_p = self.model._doc.add_paragraph(caption, style=caption_style or "Normal")
            cap_p._element.getparent().remove(cap_p._element)
            p._element.addnext(cap_p._element)

        self._rebuild()
        return True

    def replace_image(self, image_path: str, anchor: Anchor | None = None,
                      media_filename: str | None = None) -> bool:
        """Replace an existing image's binary blob. Does NOT call save()."""
        if not os.path.isfile(image_path):
            return False

        # Resolve target zip path
        zip_path = None
        if anchor and anchor.kind == "image":
            for img in self.model._images:
                if img.para_index == anchor.paragraph_index:
                    rel = self.model._doc.part.rels.get(img.r_id)
                    if rel:
                        ref = rel.target_ref
                        if ref.startswith('/'):
                            ref = ref[1:]
                        elif ref.startswith('../'):
                            ref = ref[3:]
                        if not ref.startswith('word/'):
                            ref = f'word/{ref}'
                        zip_path = ref
                    break
        elif media_filename:
            zip_path = media_filename if media_filename.startswith('word/') else f'word/{media_filename}'

        if not zip_path:
            return False

        with open(image_path, 'rb') as f:
            self._media_overrides[zip_path] = f.read()
        return True
