"""
core/creator.py — Create thesis .docx from Markdown outline.

Usage by agent:
    from core import create_thesis
    outline = '''
    # 摘要
    【摘要内容占位】
    # 第1章 绪论
    ## 1.1 研究背景
    【正文占位】
    '''
    path = create_thesis("output.docx", outline=outline)
    # Or from a template:
    path = create_thesis("output.docx", outline=outline, template="学校模板.docx")
"""

from __future__ import annotations
import os
import re
import shutil

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

from .style import CJK_THESIS_FORMAT

_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

# Default page setup for Chinese degree theses
_DEFAULT_PAGE = {
    "width_cm": 21.0,
    "height_cm": 29.7,
    "margin_top_cm": 2.54,
    "margin_bottom_cm": 2.54,
    "margin_left_cm": 3.17,
    "margin_right_cm": 3.17,
}

# Heading level → paragraph style name
_LEVEL_TO_STYLE = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
}


def create_thesis(
    output_path: str | os.PathLike,
    outline: str | None = None,
    template: str | os.PathLike | None = None,
) -> str:
    """Create a thesis .docx from a Markdown outline.

    Args:
        output_path: Where to save the new .docx.
        outline: Markdown text (or path to .md file). If None, inserts a
                 minimal default skeleton.
        template: Optional .docx template to inherit styles/page setup from.

    Returns:
        Absolute path to the created file.
    """
    output_path = os.path.abspath(output_path)

    # Read outline from file if it looks like a path
    if outline and not outline.strip().startswith("#"):
        candidate = outline.strip()
        if os.path.isfile(candidate):
            with open(candidate, encoding="utf-8") as f:
                outline = f.read()

    if template:
        return _create_from_template(output_path, outline, template)

    # Create from scratch
    doc = Document()
    _setup_page(doc)
    _ensure_styles(doc)
    _insert_outline(doc, outline or _DEFAULT_OUTLINE)
    doc.save(output_path)
    return output_path


def _create_from_template(output_path, outline, template_path):
    """Create thesis inheriting styles from a template .docx."""
    template_path = os.path.abspath(template_path)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"模板文件不存在: {template_path}")

    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    # Clear existing content paragraphs (keep styles/settings/headers)
    body = doc.element.body
    to_remove = [child for child in body if child.tag.endswith("}p") or child.tag.endswith("}tbl")]
    for child in to_remove:
        body.remove(child)

    # Build style_id map from template
    style_map = {}
    for style in doc.styles:
        if hasattr(style, "style_id"):
            style_map[style.name] = style.style_id

    _insert_outline(doc, outline or _DEFAULT_OUTLINE, style_map=style_map)
    doc.save(output_path)
    return output_path


def _setup_page(doc):
    """Set default A4 page dimensions and margins."""
    section = doc.sections[0]
    section.page_width = Cm(_DEFAULT_PAGE["width_cm"])
    section.page_height = Cm(_DEFAULT_PAGE["height_cm"])
    section.top_margin = Cm(_DEFAULT_PAGE["margin_top_cm"])
    section.bottom_margin = Cm(_DEFAULT_PAGE["margin_bottom_cm"])
    section.left_margin = Cm(_DEFAULT_PAGE["margin_left_cm"])
    section.right_margin = Cm(_DEFAULT_PAGE["margin_right_cm"])


def _ensure_styles(doc):
    """Create Heading 1/2/3 and Body Text styles if they don't exist."""
    style_formats = {
        "Heading 1": CJK_THESIS_FORMAT.get("heading_1", {}),
        "Heading 2": CJK_THESIS_FORMAT.get("heading_2", {}),
        "Heading 3": CJK_THESIS_FORMAT.get("heading_3", {}),
        "Body Text": CJK_THESIS_FORMAT.get("body", {}),
    }
    for style_name, fmt in style_formats.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            from docx.enum.style import WD_STYLE_TYPE
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        if fmt.get("font_name"):
            style.font.name = fmt["font_name"]
        if fmt.get("font_size"):
            style.font.size = Pt(fmt["font_size"])
        if fmt.get("bold") is not None:
            style.font.bold = fmt["bold"]


def _parse_markdown_outline(md_text: str) -> list[tuple[int, str]]:
    """Parse Markdown into (level, text) pairs. level=0 means body text."""
    results = []
    for line in md_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            results.append((min(level, 3), text))
        else:
            results.append((0, stripped))
    return results


def _insert_outline(doc, md_text: str, style_map: dict | None = None):
    """Insert parsed Markdown outline as styled paragraphs using python-docx API."""
    entries = _parse_markdown_outline(md_text)

    for level, text in entries:
        if level > 0:
            style_name = _LEVEL_TO_STYLE.get(level, "Normal")
        else:
            style_name = "Body Text"

        if style_map:
            # Template mode: use raw XML with template's actual style_id
            style_id = style_map.get(style_name, style_name)
            p = _make_paragraph(style_id, text)
            doc.element.body.append(p)
        else:
            # Standard mode: use python-docx API (handles style_id mapping)
            try:
                para = doc.add_paragraph(text, style=style_name)
            except KeyError:
                para = doc.add_paragraph(text)


_DEFAULT_OUTLINE = """\
# 摘要
【摘要内容占位】
关键词：大语言模型；代码生成；质量评估
# ABSTRACT
[Abstract placeholder]
Keywords: Large Language Model; Code Generation; Quality Assessment
# 第1章 绪论
## 1.1 研究背景
【正文占位】
## 1.2 研究目标
【正文占位】
# 第2章 相关工作
【正文占位】
# 第3章 方法
【正文占位】
# 第4章 实验与分析
【正文占位】
# 第5章 结论
【正文占位】
# 参考文献
[1] 【参考文献占位】
"""
