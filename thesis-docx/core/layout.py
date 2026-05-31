"""
core/layout.py — Page layout and header/footer operations.

Mixin for SafeDocument. Provides access to section-level page setup
and header/footer content.
"""

from __future__ import annotations
from typing import Optional, Any

from lxml import etree

from .types import Anchor
from .utils import cm_to_emu, emu_to_cm, W_NS, _W


class LayoutMixin:

    # ------------------------------------------------------------------
    # Page setup
    # ------------------------------------------------------------------

    def get_page_margins(self, section_index: int = 0) -> dict[str, float | None]:
        """Get margins for a section. Returns dict with top, bottom, left, right, header, footer in cm."""
        sections = self.model._doc.sections
        if section_index < 0 or section_index >= len(sections):
            return {}
        s = sections[section_index]
        return {
            "top_cm": emu_to_cm(s.top_margin),
            "bottom_cm": emu_to_cm(s.bottom_margin),
            "left_cm": emu_to_cm(s.left_margin),
            "right_cm": emu_to_cm(s.right_margin),
            "header_cm": emu_to_cm(s.header_distance),
            "footer_cm": emu_to_cm(s.footer_distance),
        }

    def set_page_margins(self, section_index: int = 0, **kwargs) -> bool:
        """Set margins for a section. Keyword args: top_cm, bottom_cm, left_cm, right_cm, header_cm, footer_cm."""
        sections = self.model._doc.sections
        if section_index < 0 or section_index >= len(sections):
            return False
        s = sections[section_index]
        for key, val_cm in kwargs.items():
            if val_cm is None:
                continue
            emu = cm_to_emu(val_cm)
            if key == "top_cm":
                s.top_margin = emu
            elif key == "bottom_cm":
                s.bottom_margin = emu
            elif key == "left_cm":
                s.left_margin = emu
            elif key == "right_cm":
                s.right_margin = emu
            elif key == "header_cm":
                s.header_distance = emu
            elif key == "footer_cm":
                s.footer_distance = emu
        return True

    def get_page_size(self, section_index: int = 0) -> dict[str, Any]:
        """Get page dimensions for a section. Returns width_cm, height_cm, orientation."""
        sections = self.model._doc.sections
        if section_index < 0 or section_index >= len(sections):
            return {}
        s = sections[section_index]
        width_cm = emu_to_cm(s.page_width)
        height_cm = emu_to_cm(s.page_height)
        orientation = "landscape" if width_cm and height_cm and width_cm > height_cm else "portrait"
        return {
            "width_cm": width_cm,
            "height_cm": height_cm,
            "orientation": orientation,
        }

    def get_sections_count(self) -> int:
        """Return the number of sections in the document."""
        return len(self.model._doc.sections)

    # ------------------------------------------------------------------
    # Headers and Footers
    # ------------------------------------------------------------------

    def get_header_text(self, section_index: int = 0) -> str | None:
        """Get the text content of the first header in a section."""
        sections = self.model._doc.sections
        if section_index < 0 or section_index >= len(sections):
            return None
        s = sections[section_index]
        try:
            header = s.header
            if header is None:
                return None
            return " ".join(p.text for p in header.paragraphs if p.text)
        except Exception:
            return None

    def set_header_text(self, text: str, section_index: int = 0,
                        alignment: str = "center") -> bool:
        """Set the text content of the first header for a section.

        Creates a header if none exists. Clears existing paragraphs first.
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        sections = self.model._doc.sections
        if section_index < 0 or section_index >= len(sections):
            return False
        s = sections[section_index]
        try:
            header = s.header
        except Exception:
            return False

        if header is None:
            return False

        # Clear existing paragraphs
        for p in header.paragraphs:
            for run in p.runs:
                run.text = ""
            p._element.getparent().remove(p._element)
        for p in list(header.paragraphs):
            try:
                p._element.getparent().remove(p._element)
            except Exception:
                pass

        # Add new paragraph
        new_p = header.add_paragraph(text)
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        if alignment in align_map:
            new_p.alignment = align_map[alignment]

        return True

    def get_footer_text(self, section_index: int = 0) -> str | None:
        """Get the text content of the first footer in a section."""
        sections = self.model._doc.sections
        if section_index < 0 or section_index >= len(sections):
            return None
        s = sections[section_index]
        try:
            footer = s.footer
            if footer is None:
                return None
            return " ".join(p.text for p in footer.paragraphs if p.text)
        except Exception:
            return None

    def set_footer_text(self, text: str, section_index: int = 0,
                        alignment: str = "center") -> bool:
        """Set the text content of the first footer for a section."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        sections = self.model._doc.sections
        if section_index < 0 or section_index >= len(sections):
            return False
        s = sections[section_index]
        try:
            footer = s.footer
        except Exception:
            return False

        if footer is None:
            return False

        for p in list(footer.paragraphs):
            try:
                p._element.getparent().remove(p._element)
            except Exception:
                pass

        new_p = footer.add_paragraph(text)
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        if alignment in align_map:
            new_p.alignment = align_map[alignment]
        return True
