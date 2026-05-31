"""
core/model.py — DocumentModel

The elegant reimplementation of the structural heart (originally ThesisDoc).

Responsibilities (strictly limited):
- Load document
- Build and maintain rich structural indexes (chapter_path is king)
- Provide powerful query capabilities via the Locator system
- Expose clean, stable data (never raw python-docx objects to upper layers for mutation)

This is intentionally *not* a god class. Mutation and persistence live elsewhere.
"""

from __future__ import annotations
import os
import re
from dataclasses import replace
from typing import Optional
from docx import Document

from .types import ParagraphInfo, SectionNode, ImageInfo, TableInfo, FormulaInfo, ReferenceInfo, Locator, Anchor
from .locator import resolve, resolve_all
from .utils import get_heading_level, get_paragraph_format, emu_to_cm

# WordprocessingML namespace for revision detection
_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_W = f'{{{_W_NS}}}'

# Math namespace for formula detection
_M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


class DocumentModel:
    """
    Clean, query-focused view of a thesis document.

    Phase 1 goal: Loading + indexes + resolve() work correctly.
    Mutation and saving are handled by SafeDocument (persistence.py).
    """

    def __init__(self, filepath: str | os.PathLike, create: bool = False):
        self.filepath = os.path.abspath(filepath)
        if create:
            self._doc = Document()
        else:
            if not os.path.exists(self.filepath):
                raise FileNotFoundError(f"文件不存在: {self.filepath}")
            self._doc = Document(self.filepath)

        self._paragraphs: list[ParagraphInfo] = []
        self._sections: list[SectionNode] = []
        self._images: list[ImageInfo] = []
        self._tables: list[TableInfo] = []
        self._formulas: list[FormulaInfo] = []
        self._references: list[ReferenceInfo] = []

        self._build_indexes()

    # ------------------------------------------------------------------
    # Public Query API (the value users actually care about)
    # ------------------------------------------------------------------

    @property
    def paragraphs(self) -> list[ParagraphInfo]:
        return self._paragraphs

    @property
    def sections(self) -> list[SectionNode]:
        return self._sections

    @property
    def images(self) -> list[ImageInfo]:
        return self._images

    @property
    def tables(self) -> list[TableInfo]:
        return self._tables

    @property
    def formulas(self) -> list[FormulaInfo]:
        return self._formulas

    @property
    def references(self) -> list[ReferenceInfo]:
        return self._references

    def resolve(self, locator: Locator | str) -> Optional[Anchor]:
        """The one true way to turn intent into a stable reference."""
        return resolve(locator, self._paragraphs, self._sections, self._images, self._tables,
                       self._formulas, self._references)

    def resolve_all(self, locator: Locator | str) -> list[Anchor]:
        """Resolve to ALL matching anchors (unlike resolve() which returns first match)."""
        return resolve_all(locator, self._paragraphs, self._sections, self._images, self._tables,
                           self._formulas, self._references)

    def find_paragraph_by_text(self, text: str, start: int = 0) -> Optional[int]:
        """Convenience method for simple text search (uses internal list for speed)."""
        for p in self._paragraphs[start:]:
            if text in p.text:
                return p.index
        return None

    def get_paragraphs_in_scope(self, scope) -> list[ParagraphInfo]:
        """Basic scope support for Phase 1 (expand later with full Locator support)."""
        if isinstance(scope, dict) and "para_range" in scope:
            start, end = scope["para_range"]
            return [p for p in self._paragraphs if start <= p.index <= end]
        return []

    # ------------------------------------------------------------------
    # Internal Index Building (kept the proven logic, improved structure)
    # ------------------------------------------------------------------

    def _build_indexes(self):
        self._build_paragraph_index()
        self._build_sections_tree()
        self._build_image_index()
        self._build_table_index()
        self._build_formula_index()
        self._build_reference_index()

    def _build_paragraph_index(self):
        self._paragraphs = []
        current_chapter_path = ""
        chapter_counters: dict[int, int] = {}

        for i, para in enumerate(self._doc.paragraphs):
            style_name = para.style.name if para.style else "Normal"
            level = get_heading_level(style_name)
            text = para.text or ""
            fmt = get_paragraph_format(para)

            if level is not None:
                chapter_counters[level] = chapter_counters.get(level, 0) + 1
                for lv in list(chapter_counters.keys()):
                    if lv > level:
                        del chapter_counters[lv]
                parts = [str(chapter_counters[lv]) for lv in sorted(chapter_counters) if lv <= level]
                current_chapter_path = ".".join(parts)

            # Revision (Track Changes) detection — Phase 1 minimal (user-accepted design)
            para_elem = para._element
            ins_count = len(para_elem.findall('.//w:ins', {'w': _W_NS}))
            del_count = len(para_elem.findall('.//w:del', {'w': _W_NS}))

            has_revisions = (ins_count + del_count) > 0
            revision_count = ins_count + del_count
            revision_types: set[str] = set()
            revision_authors: set[str] = set()

            for rev_elem in para_elem.findall('.//w:ins', {'w': _W_NS}) + para_elem.findall('.//w:del', {'w': _W_NS}):
                author = rev_elem.get(f'{_W}author')
                if author:
                    revision_authors.add(author)

                # Determine type
                if rev_elem.tag == _W + 'ins':
                    revision_types.add("insertion")
                elif rev_elem.tag == _W + 'del':
                    revision_types.add("deletion")

            info = ParagraphInfo(
                index=i,
                text=text,
                style=style_name,
                level=level,
                chapter_path=current_chapter_path,
                alignment=fmt.get("alignment"),
                first_line_indent_cm=fmt.get("first_line_indent_cm"),
                char_count=len(text),
                font_name=fmt.get("font_name"),
                font_size=fmt.get("font_size"),
                bold=fmt.get("bold"),
                italic=fmt.get("italic"),
                line_spacing=fmt.get("line_spacing"),
                line_spacing_rule=fmt.get("line_spacing_rule"),
                space_before=fmt.get("space_before"),
                space_after=fmt.get("space_after"),
                has_revisions=has_revisions,
                revision_count=revision_count,
                revision_types=frozenset(revision_types),
                revision_authors=frozenset(revision_authors),
            )
            self._paragraphs.append(info)

    def _build_sections_tree(self):
        """Build hierarchical section tree with correct para_range (adapted cleanly from original)."""
        heading_paras = [p for p in self._paragraphs if p.level is not None]
        if not heading_paras:
            self._sections = []
            return

        total = len(self._paragraphs)

        def build_node(idx: int, parent_path: str = ""):
            hp = heading_paras[idx]
            level = hp.level
            path = hp.chapter_path
            para_start = hp.index
            para_end = total - 1

            for j in range(idx + 1, len(heading_paras)):
                if heading_paras[j].level <= level:
                    para_end = heading_paras[j].index - 1
                    break

            children = []
            ci = idx + 1
            while ci < len(heading_paras) and heading_paras[ci].level > level:
                if heading_paras[ci].level == level + 1:
                    child, ci = build_node(ci, path)
                    children.append(child)
                else:
                    ci += 1

            char_count = sum(
                self._paragraphs[p].char_count
                for p in range(para_start, min(para_end + 1, total))
            )

            # Aggregate revision stats for this section (Phase 1 enhancement)
            # Bottom-up: from children + direct paragraphs in range
            rev_count = 0
            rev_types: set[str] = set()
            for child in children:
                if child.has_revisions:
                    rev_count += child.revision_count
                    rev_types.update(child.revision_types)

            for p_idx in range(para_start, min(para_end + 1, total)):
                p = self._paragraphs[p_idx]
                if p.has_revisions:
                    rev_count += p.revision_count
                    rev_types.update(p.revision_types)

            has_revs = rev_count > 0

            node = SectionNode(
                level=level,
                title=hp.text,
                para_range=(para_start, para_end),
                chapter_path=path,
                children=children,
                has_revisions=has_revs,
                revision_count=rev_count,
                revision_types=frozenset(rev_types),
            )
            return node, ci

        self._sections = []
        i = 0
        while i < len(heading_paras):
            hp = heading_paras[i]
            if hp.level == 1:
                node, i = build_node(i)
                self._sections.append(node)
            else:
                i += 1

    # ------------------------------------------------------------------
    # Image and Table Indexing - E1 Improved Version
    # ------------------------------------------------------------------

    def _find_nearby_caption(self, para_index: int, search_range: int = 4) -> Optional[str]:
        """
        E1 improved caption finder.
        Priority:
        1. Paragraph with style containing 'caption' (or 图注/表注).
        2. Text matching standard thesis numbering (图 2-1, 表3.2, Figure 1-3, etc.).
        Prefers closer matches and exact style matches.
        """
        caption_pattern = re.compile(r'^\s*(图|表|Figure|Table|Fig\.?)\s*\d+[\.\-－]\d+')
        style_caption_keywords = ("caption", "图注", "表注", "插图", "表格")

        best_caption = None
        best_score = -1

        for offset in range(-search_range, search_range + 1):
            idx = para_index + offset
            if not (0 <= idx < len(self._paragraphs)):
                continue

            p = self._paragraphs[idx]
            text = p.text.strip()
            if not text or len(text) > 120:
                continue

            score = 0

            # Strong signal: caption style
            if p.style:
                style_lower = p.style.lower()
                if any(kw in style_lower for kw in style_caption_keywords):
                    score += 10

            # Good signal: standard numbering pattern
            if caption_pattern.match(text):
                score += 6

            # Bonus for being very close
            distance = abs(offset)
            score += max(0, 4 - distance)

            if score > best_score:
                best_score = score
                best_caption = text

        return best_caption

    def _build_image_index(self):
        """E1 version: Better caption detection + chapter_path."""
        self._images = []
        image_rels = {}
        for rel in self._doc.part.rels.values():
            if "image" in rel.reltype:
                image_rels[rel.rId] = rel

        for i, para in enumerate(self._doc.paragraphs):
            drawings = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            for drawing in drawings:
                blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                if blip is not None:
                    embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed and embed in image_rels:
                        rel = image_rels[embed]
                        width_cm = None
                        height_cm = None
                        extent = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                        if extent is not None:
                            cx = extent.get('cx')
                            cy = extent.get('cy')
                            if cx and cy:
                                width_cm = emu_to_cm(int(cx))
                                height_cm = emu_to_cm(int(cy))

                        caption = self._find_nearby_caption(i)

                        ch_path = self._paragraphs[i].chapter_path if i < len(self._paragraphs) else ""
                        img_info = ImageInfo(
                            para_index=i,
                            r_id=embed,
                            filename=rel.target_ref,
                            format=os.path.splitext(rel.target_ref)[1].lstrip("."),
                            caption=caption,
                            width_cm=width_cm,
                            height_cm=height_cm,
                            chapter_path=ch_path,
                        )
                        self._images.append(img_info)

        # Backfill has_image on ParagraphInfo
        for img in self._images:
            idx = img.para_index
            if idx < len(self._paragraphs) and not self._paragraphs[idx].has_image:
                p = self._paragraphs[idx]
                self._paragraphs[idx] = replace(p, has_image=True)

    def _build_table_index(self):
        """E1 version: XML body walk for accurate para_index + shared caption finder."""
        self._tables = []

        body = self._doc.element.body
        _W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        p_tag = f'{{{_W_NS}}}p'

        for t_idx, table in enumerate(self._doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            header = rows[0] if rows else []
            row_count = len(rows)
            col_count = len(header) if header else (len(rows[0]) if rows else 0)

            # Walk body XML to count paragraphs before this table element
            tbl_element = table._element
            para_index = 0
            for child in body:
                if child is tbl_element:
                    break
                if child.tag == p_tag:
                    para_index += 1

            caption = self._find_nearby_caption(para_index, search_range=5)

            # Try to find the actual paragraph that has this exact caption text
            best_para_idx = para_index
            if caption:
                for offset in range(-5, 6):
                    idx = para_index + offset
                    if 0 <= idx < len(self._paragraphs) and self._paragraphs[idx].text.strip() == caption:
                        best_para_idx = idx
                        break

            ch_path = self._paragraphs[best_para_idx].chapter_path if best_para_idx < len(self._paragraphs) else ""

            self._tables.append(TableInfo(
                para_index=best_para_idx,
                caption=caption,
                header=header,
                chapter_path=ch_path,
                row_count=row_count,
                col_count=col_count,
            ))

    # ------------------------------------------------------------------
    # Formula Index
    # ------------------------------------------------------------------

    def _build_formula_index(self):
        """Scan paragraphs for OMML/OLE/placeholder formulas."""
        self._formulas = []
        m_ns = f'{{{_M_NS}}}'
        eq_pattern = re.compile(r'\((\d+[\.\-]\d+|\d+)\)')

        for i, para in enumerate(self._doc.paragraphs):
            elem = para._element

            formula_type = None
            content = ""

            # Check for OMML formulas
            omath_elems = elem.findall(f'.//{m_ns}oMath')
            if omath_elems:
                formula_type = "OMML"
                parts = []
                for omath in omath_elems:
                    for mr in omath.findall(f'.//{m_ns}r'):
                        mt = mr.find(f'{m_ns}t')
                        if mt is not None and mt.text:
                            parts.append(mt.text)
                content = ''.join(parts)

            # Check for OLE objects
            if not formula_type:
                if elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object'):
                    formula_type = "OLE"

            if not formula_type:
                continue

            # Extract equation number from paragraph text
            eq_match = eq_pattern.search(para.text or "")
            eq_number = eq_match.group(1) if eq_match else None

            ch_path = self._paragraphs[i].chapter_path if i < len(self._paragraphs) else ""

            self._formulas.append(FormulaInfo(
                para_index=i,
                formula_type=formula_type,
                content=content,
                equation_number=eq_number,
                chapter_path=ch_path,
            ))

    # ------------------------------------------------------------------
    # Reference Index
    # ------------------------------------------------------------------

    def _build_reference_index(self):
        """Find and parse the bibliography section.

        Handles two numbering styles:
        1. Explicit: text starts with [N] (e.g. "[1] Author. Title...")
        2. Auto-numbered: List Paragraph style with w:numPr (number rendered by Word, not in text)
        """
        self._references = []
        ref_num_pattern = re.compile(r'^\[(\d+)\]\s*')
        type_markers = {'[J]': 'journal', '[C]': 'conference', '[M]': 'book',
                        '[D]': 'thesis', '[EB': 'online', '[N]': 'newspaper'}

        # Find the reference section
        ref_start = None
        ref_chapter_path = ""
        for p in self._paragraphs:
            if p.level is not None and p.level <= 2 and '参考文献' in p.text:
                ref_start = p.index
                ref_chapter_path = p.chapter_path
                break

        if ref_start is None:
            return

        # Collect candidate paragraphs until next heading of same or higher level
        candidates = []
        for i in range(ref_start + 1, len(self._paragraphs)):
            p = self._paragraphs[i]
            if p.level is not None and p.level <= 2:
                break
            if not p.text.strip():
                continue
            candidates.append(i)

        if not candidates:
            return

        # First pass: try explicit [N] numbering
        explicit_found = False
        for i in candidates:
            p = self._paragraphs[i]
            match = ref_num_pattern.match(p.text)
            if match:
                explicit_found = True
                break

        if explicit_found:
            # Only include paragraphs with explicit [N] prefix
            seq = 1
            for i in candidates:
                p = self._paragraphs[i]
                match = ref_num_pattern.match(p.text)
                if not match:
                    continue
                num = int(match.group(1))
                text = p.text[match.end():]
                ref_type = None
                for marker, rtype in type_markers.items():
                    if marker in text:
                        ref_type = rtype
                        break
                self._references.append(ReferenceInfo(
                    index=num, para_index=i, text=text,
                    ref_type=ref_type, chapter_path=ref_chapter_path,
                ))
        else:
            # Auto-numbered: all non-empty paragraphs in the section are references.
            # Assign sequential numbers.
            for seq, i in enumerate(candidates, start=1):
                p = self._paragraphs[i]
                ref_type = None
                for marker, rtype in type_markers.items():
                    if marker in p.text:
                        ref_type = rtype
                        break
                self._references.append(ReferenceInfo(
                    index=seq, para_index=i, text=p.text,
                    ref_type=ref_type, chapter_path=ref_chapter_path,
                ))
