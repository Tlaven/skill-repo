"""
core/style.py — Style and format operations.

Mixin for SafeDocument. Provides style assignment, detection, clearing,
heading auto-classification, format fixers, and detect→fix→re-detect loop
for Chinese academic theses.
"""

from __future__ import annotations
import re
from dataclasses import replace
from typing import Optional, Any

from lxml import etree

from .types import Anchor, ParagraphInfo
from .utils import W_NS, _W, emu_to_cm, cm_to_emu, pt_to_emu


# =============================================================================
# Chinese Academic Thesis Format Standards
# =============================================================================

# Standard format specifications for Chinese degree theses.
# Used by fix_format() and detect_format_issues().
# Values can be overridden by passing a custom config dict.

CJK_THESIS_FORMAT: dict[str, dict[str, Any]] = {
    "body": {
        "font_name": "宋体",
        "font_name_ascii": "Times New Roman",
        "font_name_east_asia": "宋体",
        "font_size": 12.0,            # 小四
        "bold": False,
        "italic": False,
        "alignment": "justify",
        "line_spacing": 1.5,
        "line_spacing_rule": "multiple",
        "first_line_indent_cm": 0.74,  # ≈2 汉字字符
        "space_before": 0,
        "space_after": 0,
    },
    "heading_1": {
        "font_name": "黑体",
        "font_name_ascii": "Times New Roman",
        "font_size": 16.0,            # 小三
        "bold": True,
        "italic": False,
        "alignment": "center",
        "line_spacing": 1.25,
        "line_spacing_rule": "multiple",
        "first_line_indent_cm": None,
        "space_before": 12.0,
        "space_after": 6.0,
    },
    "heading_2": {
        "font_name": "黑体",
        "font_name_ascii": "Times New Roman",
        "font_size": 14.0,            # 四号
        "bold": True,
        "italic": False,
        "alignment": "left",
        "line_spacing": 1.25,
        "line_spacing_rule": "multiple",
        "first_line_indent_cm": None,
        "space_before": 6.0,
        "space_after": 3.0,
    },
    "heading_3": {
        "font_name": "黑体",
        "font_name_ascii": "Times New Roman",
        "font_size": 13.0,            # 小四
        "bold": True,
        "italic": False,
        "alignment": "left",
        "line_spacing": 1.25,
        "line_spacing_rule": "multiple",
        "first_line_indent_cm": None,
        "space_before": 3.0,
        "space_after": 3.0,
    },
    "caption": {
        "font_name": "宋体",
        "font_name_ascii": "Times New Roman",
        "font_size": 10.5,            # 五号
        "bold": False,
        "italic": False,
        "alignment": "center",
        "line_spacing": 1.0,
        "line_spacing_rule": "multiple",
        "first_line_indent_cm": None,
        "space_before": 3.0,
        "space_after": 3.0,
    },
    "reference": {
        "font_name": "宋体",
        "font_name_ascii": "Times New Roman",
        "font_size": 10.5,            # 五号
        "bold": False,
        "italic": False,
        "alignment": "justify",
        "line_spacing": 1.25,
        "line_spacing_rule": "multiple",
        "first_line_indent_cm": None,
        "space_before": 0,
        "space_after": 0,
    },
}

# Paragraph role → format key mapping
ROLE_MAP = {
    "body": "body",
    "heading_1": "heading_1",
    "heading_2": "heading_2",
    "heading_3": "heading_3",
    "caption": "caption",
    "reference": "reference",
}


class StyleMixin:

    def set_style(self, anchor: Anchor, style_name: str) -> bool:
        """Assign a named style to the paragraph referenced by anchor."""
        if anchor.kind != "paragraph" or anchor.paragraph_index is None:
            return False
        idx = anchor.paragraph_index
        if idx < 0 or idx >= len(self.model._doc.paragraphs):
            return False
        try:
            para = self.model._doc.paragraphs[idx]
            para.style = self.model._doc.styles[style_name]
            if idx < len(self.model._paragraphs):
                p = self.model._paragraphs[idx]
                self.model._paragraphs[idx] = replace(p, style=style_name)
            return True
        except KeyError:
            return False

    def get_style_names(self) -> list[str]:
        return [s.name for s in self.model._doc.styles]

    def detect_style_usage(self) -> dict[str, int]:
        counts: dict[str, int] = {}
        for p in self.model._paragraphs:
            counts[p.style] = counts.get(p.style, 0) + 1
        counts["_guide"] = "当前文档样式使用统计。正文/标题/图题应各用什么样式，标准见 references/format-roles.md。"
        return counts

    def clear_direct_formatting(self, anchor: Anchor) -> bool:
        """Remove inline run formatting from a paragraph."""
        if anchor.kind != "paragraph" or anchor.paragraph_index is None:
            return False
        idx = anchor.paragraph_index
        if idx < 0 or idx >= len(self.model._doc.paragraphs):
            return False
        para = self.model._doc.paragraphs[idx]
        for run in para.runs:
            rPr = run._element.find(f'{_W}rPr')
            if rPr is not None:
                run._element.remove(rPr)
        return True

    def assign_headings(self, dry_run: bool = False) -> dict:
        """Auto-detect heading paragraphs and assign heading styles."""
        stats = {"assigned": 0, "skipped": 0, "errors": [], "candidates": []}
        heading_1_pattern = re.compile(r'^\s*第[一二三四五六七八九十百\d]+章')
        heading_2_pattern = re.compile(r'^\s*[（(]?[一二三四五六七八九十百]+[）).、]\s*')
        outline_pattern = re.compile(r'^\s*(\d+)(?:\.(\d+))?(?:\.(\d+))?\s+')

        candidates = []
        for i, p in enumerate(self.model._paragraphs):
            if p.level is not None:
                stats["skipped"] += 1
                continue
            text = p.text.strip()
            if not text:
                continue
            level = None
            if heading_1_pattern.match(text):
                level = 1
            elif outline_pattern.match(text):
                m = outline_pattern.match(text)
                groups = [g for g in m.groups() if g]
                level = min(len(groups), 3)
            elif heading_2_pattern.match(text):
                level = 2
            if level is not None and len(text) <= 60:
                candidates.append((i, level))

        stats["candidates"] = len(candidates)
        if dry_run:
            stats["assigned"] = len(candidates)
            return stats

        for idx, level in candidates:
            style_name = f"Heading {level}"
            try:
                para = self.model._doc.paragraphs[idx]
                para.style = self.model._doc.styles[style_name]
                p = self.model._paragraphs[idx]
                self.model._paragraphs[idx] = replace(p, style=style_name, level=level)
                stats["assigned"] += 1
            except KeyError:
                stats["errors"].append(f"para {idx}: style '{style_name}' not found")
        return stats

    def detect_format_issues(self, with_guide: bool = False) -> list[dict]:
        """Scan document for common formatting issues."""
        issues = []
        styles = self.get_style_names()

        for p in self.model._paragraphs:
            idx = p.index

            if p.level is None and p.font_size is not None:
                if p.font_size < 10:
                    issues.append(self._make_issue(idx, p, "font_too_small",
                                                   f"Size {p.font_size}pt < 10pt"))
                elif p.font_size > 16:
                    issues.append(self._make_issue(idx, p, "font_too_large",
                                                   f"Size {p.font_size}pt > 16pt"))

            COMMON_FONTS = {"宋体", "SimSun", "Times New Roman", "楷体", "KaiTi",
                            "仿宋_GB2312", "FangSong_GB2312", "黑体", "SimHei",
                            "Arial", "Cambria Math"}
            if p.level is None and p.font_name:
                if p.font_name not in COMMON_FONTS:
                    issues.append(self._make_issue(idx, p, "unusual_font",
                                                   f"Font '{p.font_name}'"))

            if p.style and p.style not in styles:
                issues.append(self._make_issue(idx, p, "missing_style_definition",
                                               f"Style '{p.style}' not in document"))

        # Check heading hierarchy
        prev_level = 0
        for p in self.model._paragraphs:
            if p.level is not None:
                if p.level > prev_level + 1 and prev_level > 0:
                    issues.append(self._make_issue(p.index, p, "heading_hierarchy",
                                                   f"Jumps from level {prev_level} to {p.level}"))
                prev_level = p.level

        # Check images for caption
        for img in self.model._images:
            if not img.caption:
                issues.append({
                    "severity": "warning",
                    "para_index": img.para_index,
                    "chapter": img.chapter_path,
                    "text_snippet": img.filename,
                    "issue": "image_no_caption",
                    "detail": f"Image '{img.filename}' has no caption",
                })

        # Check tables for caption
        for t in self.model._tables:
            if not t.caption:
                issues.append({
                    "severity": "warning",
                    "para_index": t.para_index,
                    "chapter": t.chapter_path,
                    "text_snippet": f"Table at para {t.para_index}",
                    "issue": "table_no_caption",
                    "detail": "Table has no caption",
                })

        issues.sort(key=lambda x: (x.get("para_index", 0) or 0))
        if with_guide:
            issues.append({
                "_guide": True,
                "message": "以上是检测到的格式问题。各项正确标准见 references/format-roles.md；自动修复用 fix_format()。",
                "file": "references/format-roles.md",
            })
        return issues

    def _make_issue(self, idx, p, issue, detail, severity="info"):
        return {
            "severity": severity,
            "para_index": idx,
            "chapter": p.chapter_path if hasattr(p, 'chapter_path') else getattr(p, 'chapter_path', ''),
            "text_snippet": getattr(p, 'text', '')[:50],
            "issue": issue,
            "detail": detail,
        }

    # ==================================================================
    # Writing Quality Checks (extract data for LLM judgment)
    # ==================================================================

    def check_writing_style(self) -> dict:
        body_paras = [
            {
                "index": p.index,
                "chapter": p.chapter_path,
                "text": p.text[:200],
            }
            for p in self.model._paragraphs
            if p.level is None and p.char_count > 30
        ]
        return {
            "body_paragraphs": len(body_paras),
            "samples": body_paras[:20],
            "_guide": "检查套话、机械过渡、空泛结论。判断标准见 references/cliches.md。",
        }

    def check_terminology(self) -> dict:
        import re
        abbrev_pattern = re.compile(r'\b[A-Z]{2,6}\b')
        cn_en_pattern = re.compile(r'[\u4e00-\u9fff]+[A-Za-z]+|[\u4e00-\u9fff]+\([A-Za-z]+\)')
        seen_abbrevs = set()
        seen_cn_en = set()
        for p in self.model._paragraphs:
            if p.level is not None:
                continue
            for m in abbrev_pattern.finditer(p.text):
                seen_abbrevs.add(m.group())
            for m in cn_en_pattern.finditer(p.text):
                candidate = m.group()[:40]
                if len(candidate) > 4:
                    seen_cn_en.add(candidate)
        return {
            "abbreviations": sorted(seen_abbrevs),
            "cn_en_terms": sorted(seen_cn_en)[:30],
            "_guide": "疑似术语候选列表。检查这些术语是否全文一致、缩写是否首次定义。标准见 references/terminology.md。",
        }

    # ==================================================================
    # Format Setters
    # ==================================================================

    def set_run_property(self, anchor: Anchor, **kwargs) -> bool:
        """Set run-level formatting properties.

        Kwargs: font_name, font_name_ascii, font_name_east_asia,
                font_size (pt), bold, italic.
        Overwrites the entire rPr of all runs in the paragraph.
        """
        if anchor.kind != "paragraph" or anchor.paragraph_index is None:
            return False
        idx = anchor.paragraph_index
        if idx < 0 or idx >= len(self.model._doc.paragraphs):
            return False
        if not kwargs:
            return False

        para = self.model._doc.paragraphs[idx]
        for run in para.runs:
            rPr = run._element.find(f'{_W}rPr')
            if rPr is None:
                rPr = etree.SubElement(run._element, f'{_W}rPr')
                run._element.insert(0, rPr)

            for key, val in kwargs.items():
                if val is None:
                    continue
                if key == "font_name":
                    self._set_rpr_font(rPr, val)
                elif key == "font_name_ascii":
                    self._set_rpr_font_ascii(rPr, val)
                elif key == "font_name_east_asia":
                    self._set_rpr_font_east_asia(rPr, val)
                elif key == "font_size":
                    self._set_rpr_size(rPr, val)
                elif key == "bold":
                    self._set_rpr_bold_italic(rPr, 'b', val)
                elif key == "italic":
                    self._set_rpr_bold_italic(rPr, 'i', val)

        return True

    def _set_rpr_font(self, rPr, font_name):
        rFonts = rPr.find(f'{_W}rFonts')
        if rFonts is None:
            rFonts = etree.SubElement(rPr, f'{_W}rFonts')
        rFonts.set(f'{_W}ascii', font_name)
        rFonts.set(f'{_W}hAnsi', font_name)
        rFonts.set(f'{_W}cs', font_name)

    def _set_rpr_font_ascii(self, rPr, font_name):
        rFonts = rPr.find(f'{_W}rFonts')
        if rFonts is None:
            rFonts = etree.SubElement(rPr, f'{_W}rFonts')
        rFonts.set(f'{_W}ascii', font_name)
        rFonts.set(f'{_W}hAnsi', font_name)

    def _set_rpr_font_east_asia(self, rPr, font_name):
        rFonts = rPr.find(f'{_W}rFonts')
        if rFonts is None:
            rFonts = etree.SubElement(rPr, f'{_W}rFonts')
        rFonts.set(f'{_W}eastAsia', font_name)

    def _set_rpr_size(self, rPr, size_pt):
        val = str(int(size_pt * 2))
        sz = rPr.find(f'{_W}sz')
        if sz is None:
            sz = etree.SubElement(rPr, f'{_W}sz')
        sz.set(f'{_W}val', val)
        szCs = rPr.find(f'{_W}szCs')
        if szCs is None:
            szCs = etree.SubElement(rPr, f'{_W}szCs')
        szCs.set(f'{_W}val', val)

    def _set_rpr_bold_italic(self, rPr, tag, value):
        elem = rPr.find(f'{_W}{tag}')
        if value:
            if elem is None:
                elem = etree.SubElement(rPr, f'{_W}{tag}')
            if f'{_W}val' in (elem.attrib or {}):
                elem.set(f'{_W}val', 'true')
        else:
            if elem is None:
                elem = etree.SubElement(rPr, f'{_W}{tag}')
            elem.set(f'{_W}val', '0')

    def set_alignment(self, anchor: Anchor, alignment: str) -> bool:
        """Set paragraph alignment. alignment: left, center, right, justify."""
        if anchor.kind != "paragraph" or anchor.paragraph_index is None:
            return False
        idx = anchor.paragraph_index
        if idx < 0 or idx >= len(self.model._doc.paragraphs):
            return False
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if alignment not in align_map:
            return False
        para = self.model._doc.paragraphs[idx]
        para.alignment = align_map[alignment]
        return True

    def set_line_spacing(self, anchor: Anchor, spacing: float,
                         rule: str = "multiple") -> bool:
        """Set paragraph line spacing.

        rule: 'multiple' (1.0=单倍, 1.5=一倍半),
              'exact' (spacing in pt),
              'atLeast' (minimum spacing in pt).
        """
        if anchor.kind != "paragraph" or anchor.paragraph_index is None:
            return False
        idx = anchor.paragraph_index
        if idx < 0 or idx >= len(self.model._doc.paragraphs):
            return False
        para = self.model._doc.paragraphs[idx]
        pf = para._element.find(f'{_W}pPr')
        if pf is None:
            pf = etree.SubElement(para._element, f'{_W}pPr')
            para._element.insert(0, pf)
        spacing_elem = pf.find(f'{_W}spacing')
        if spacing_elem is None:
            spacing_elem = etree.SubElement(pf, f'{_W}spacing')

        if rule == "multiple":
            val = int(spacing * 240)
            spacing_elem.set(f'{_W}line', str(val))
            spacing_elem.set(f'{_W}lineRule', 'auto')
        elif rule == "exact":
            val = int(spacing * 20)
            spacing_elem.set(f'{_W}line', str(val))
            spacing_elem.set(f'{_W}lineRule', 'exact')
        elif rule == "atLeast":
            val = int(spacing * 20)
            spacing_elem.set(f'{_W}line', str(val))
            spacing_elem.set(f'{_W}lineRule', 'atLeast')
        else:
            return False
        return True

    def set_indent(self, anchor: Anchor, first_line_cm: float | None = None,
                   left_cm: float | None = None) -> bool:
        """Set paragraph indentation."""
        if anchor.kind != "paragraph" or anchor.paragraph_index is None:
            return False
        idx = anchor.paragraph_index
        if idx < 0 or idx >= len(self.model._doc.paragraphs):
            return False
        para = self.model._doc.paragraphs[idx]
        pf = para._element.find(f'{_W}pPr')
        if pf is None:
            pf = etree.SubElement(para._element, f'{_W}pPr')
            para._element.insert(0, pf)
        ind = pf.find(f'{_W}ind')
        if ind is None:
            ind = etree.SubElement(pf, f'{_W}ind')

        if first_line_cm is not None:
            ind.set(f'{_W}firstLine', str(int(first_line_cm * 567)))
        if left_cm is not None:
            ind.set(f'{_W}left', str(int(left_cm * 567)))
        return True

    def set_space_before_after(self, anchor: Anchor,
                               before_pt: float | None = None,
                               after_pt: float | None = None) -> bool:
        """Set paragraph spacing before/after in points."""
        if anchor.kind != "paragraph" or anchor.paragraph_index is None:
            return False
        idx = anchor.paragraph_index
        if idx < 0 or idx >= len(self.model._doc.paragraphs):
            return False
        para = self.model._doc.paragraphs[idx]
        pf = para._element.find(f'{_W}pPr')
        if pf is None:
            pf = etree.SubElement(para._element, f'{_W}pPr')
            para._element.insert(0, pf)
        spacing_elem = pf.find(f'{_W}spacing')
        if spacing_elem is None:
            spacing_elem = etree.SubElement(pf, f'{_W}spacing')
        if before_pt is not None:
            spacing_elem.set(f'{_W}before', str(int(before_pt * 20)))
        if after_pt is not None:
            spacing_elem.set(f'{_W}after', str(int(after_pt * 20)))
        return True

    def set_page_orientation(self, section_index: int = 0,
                             orientation: str = "portrait") -> bool:
        """Set page orientation. orientation: portrait, landscape."""
        sections = self.model._doc.sections
        if section_index < 0 or section_index >= len(sections):
            return False
        s = sections[section_index]
        width = s.page_width
        height = s.page_height
        if orientation == "landscape":
            if s.page_width < s.page_height:
                s.page_width = height
                s.page_height = width
        elif orientation == "portrait":
            if s.page_width > s.page_height:
                s.page_width = height
                s.page_height = width
        else:
            return False
        return True

    # ==================================================================
    # Detect role for a paragraph
    # ==================================================================

    def _detect_paragraph_role(self, p: ParagraphInfo) -> str:
        """Determine formatting role for a paragraph.

        Returns one of: heading_1, heading_2, heading_3, caption, reference, body.
        """
        if p.level == 1:
            return "heading_1"
        if p.level == 2:
            return "heading_2"
        if p.level == 3:
            return "heading_3"

        text = p.text.strip()

        # Caption detection (图/表 prefix)
        if re.match(r'^\s*(图|表|Figure|Table|Fig\.?)\s*\d+[\.\-－]', text):
            return "caption"

        # Reference section detection
        ref_start = None
        for pp in self.model._paragraphs:
            if pp.level is not None and pp.level <= 2 and '参考文献' in pp.text:
                ref_start = pp.index
                break
        if ref_start is not None and p.index > ref_start:
            if p.level is None:
                return "reference"

        return "body"

    # ==================================================================
    # Fix Format — auto scan and repair
    # ==================================================================

    def fix_format(self, config: dict | None = None,
                   roles: set[str] | None = None,
                   dry_run: bool = False) -> dict:
        """Auto-format all paragraphs according to thesis standards.

        Args:
            config: Format spec dict (defaults to CJK_THESIS_FORMAT).
            roles: Which roles to fix, e.g. {"body", "heading_1"}.
                   None = all roles.
            dry_run: Report only, don't modify.

        Returns: stats dict with changes per role.
        """
        cfg = config or CJK_THESIS_FORMAT
        stats: dict[str, Any] = {
            "total_scanned": 0,
            "total_changed": 0,
            "by_role": {},
            "details": [],
        }

        for p in self.model._paragraphs:
            role = self._detect_paragraph_role(p)
            if role not in cfg:
                continue
            if roles is not None and role not in roles:
                continue

            fmt = cfg[role]
            fmt.setdefault("font_name_ascii", "Times New Roman")
            fmt.setdefault("font_name_east_asia", fmt.get("font_name", "宋体"))
            stats["total_scanned"] += 1

            changes = []
            a = Anchor(kind="paragraph", paragraph_index=p.index,
                       text_snippet=p.text[:80], chapter_path=p.chapter_path)

            # Run properties
            run_props = {}
            for prop in ("font_name", "font_name_ascii", "font_name_east_asia",
                         "font_size", "bold", "italic"):
                val = fmt.get(prop)
                if val is not None:
                    old = getattr(p, prop, None)
                    if old != val:
                        run_props[prop] = val
                        changes.append(f"{prop}: {old}→{val}")

            if run_props and not dry_run:
                self.set_run_property(a, **run_props)

            # Alignment
            align_val = fmt.get("alignment")
            if align_val is not None and p.alignment != align_val and not dry_run:
                self.set_alignment(a, align_val)
                changes.append(f"align: {p.alignment}→{align_val}")

            # Line spacing
            ls_val = fmt.get("line_spacing")
            ls_rule = fmt.get("line_spacing_rule", "multiple")
            if ls_val is not None:
                old_ls = p.line_spacing
                if old_ls != ls_val and not dry_run:
                    self.set_line_spacing(a, ls_val, ls_rule)
                    changes.append(f"line_spacing: {old_ls}→{ls_val}")

            # First line indent
            fli_val = fmt.get("first_line_indent_cm")
            if fli_val is not None:
                old_fli = p.first_line_indent_cm
                if old_fli != fli_val and not dry_run:
                    self.set_indent(a, first_line_cm=fli_val)
                    changes.append(f"indent: {old_fli}→{fli_val}")
            elif fmt.get("first_line_indent_cm") is None:
                pass  # heading/caption: no indent, already set

            # Space before/after
            sb_val = fmt.get("space_before")
            sa_val = fmt.get("space_after")
            if (sb_val is not None or sa_val is not None) and not dry_run:
                old_sb = p.space_before
                old_sa = p.space_after
                if old_sb != sb_val or old_sa != sa_val:
                    self.set_space_before_after(a, before_pt=sb_val, after_pt=sa_val)
                    changes.append(f"space: ({old_sb},{old_sa})→({sb_val},{sa_val})")

            if changes:
                stats["total_changed"] += 1
                entry = {"para_index": p.index, "role": role, "changes": changes}
                stats["by_role"].setdefault(role, 0)
                stats["by_role"][role] += 1
                stats["details"].append(entry)

        return stats

    def detect_and_fix(self, config: dict | None = None,
                       max_iterations: int = 3,
                       roles: set[str] | None = None) -> dict:
        """Run detect → fix → re-detect loop until clean or max_iterations.

        Returns final report with per-iteration stats.
        """
        report = {
            "iterations": [],
            "all_clean": False,
        }

        for iteration in range(1, max_iterations + 1):
            issues = self.detect_format_issues()
            iter_report = {
                "iteration": iteration,
                "issues_before": len(issues),
                "fix_stats": None,
                "issues_after": None,
            }

            if not issues:
                iter_report["issues_after"] = 0
                report["iterations"].append(iter_report)
                report["all_clean"] = True
                break

            fix_stats = self.fix_format(config=config, roles=roles)
            iter_report["fix_stats"] = fix_stats

            issues_after = self.detect_format_issues()
            iter_report["issues_after"] = len(issues_after)

            report["iterations"].append(iter_report)

            if not issues_after:
                report["all_clean"] = True
                break

        return report
