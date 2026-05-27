"""ThesisDoc 核心类 — 文档加载、索引构建、保存"""
import os
import re
from docx import Document
from lib.utils import (
    get_heading_level, get_run_font_info, get_paragraph_format,
    emu_to_cm,
)


class ThesisDoc:
    """论文文档的核心抽象，所有模块的入口"""

    def __init__(self, filepath, create=False):
        self.filepath = os.path.abspath(filepath)
        if create:
            self.doc = Document()
        else:
            if not os.path.exists(self.filepath):
                raise FileNotFoundError(f"文件不存在: {self.filepath}")
            self.doc = Document(self.filepath)
        self._para_index = []
        self._sections = []
        self._images = []
        self._tables = []
        self._build_index()

    def _build_index(self):
        self._build_para_index()
        self._build_sections_tree()
        self._build_image_index()
        self._build_table_index()

    def _build_para_index(self):
        self._para_index = []
        current_chapter_path = ""
        chapter_counters = {}
        for i, para in enumerate(self.doc.paragraphs):
            style_name = para.style.name if para.style else "Normal"
            level = get_heading_level(style_name)
            text = para.text or ""
            fmt = get_paragraph_format(para)
            if level is not None:
                chapter_counters[level] = chapter_counters.get(level, 0) + 1
                for lv in list(chapter_counters.keys()):
                    if lv > level:
                        del chapter_counters[lv]
                parts = []
                for lv in sorted(chapter_counters.keys()):
                    if lv <= level:
                        parts.append(str(chapter_counters[lv]))
                current_chapter_path = ".".join(parts)
            has_image, image_ids = self._check_para_images(para)
            info = {
                "index": i,
                "text": text,
                "style": style_name,
                "level": level,
                "alignment": fmt["alignment"],
                "line_spacing": fmt["line_spacing"],
                "first_line_indent": fmt["first_line_indent_cm"],
                "runs": [get_run_font_info(r) for r in para.runs],
                "has_image": has_image,
                "image_ids": image_ids,
                "chapter_path": current_chapter_path if level is None else current_chapter_path,
                "char_count": len(text),
            }
            if level is not None:
                info["chapter_path"] = current_chapter_path
            self._para_index.append(info)

    def _check_para_images(self, para):
        image_ids = []
        drawings = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        for drawing in drawings:
            blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
            if blip is not None:
                embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed:
                    image_ids.append(embed)
        if not image_ids:
            inline = para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
            for inl in inline:
                blip = inl.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                if blip is not None:
                    embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed:
                        image_ids.append(embed)
        return len(image_ids) > 0, image_ids

    def _build_sections_tree(self):
        heading_paras = []
        for info in self._para_index:
            if info["level"] is not None:
                heading_paras.append(info)
        if not heading_paras:
            return
        total = len(self._para_index)

        def build_node_single(idx, parent_path=""):
            hp = heading_paras[idx]
            level = hp["level"]
            path = hp["chapter_path"]
            para_start = hp["index"]
            para_end = total - 1
            for j in range(idx + 1, len(heading_paras)):
                if heading_paras[j]["level"] <= level:
                    para_end = heading_paras[j]["index"] - 1
                    break
            children = []
            ci = idx + 1
            while ci < len(heading_paras) and heading_paras[ci]["level"] > level:
                if heading_paras[ci]["level"] == level + 1:
                    child, ci = build_node_single(ci, path)
                    children.append(child)
                else:
                    ci += 1
            char_count = sum(
                self._para_index[p]["char_count"]
                for p in range(para_start, min(para_end + 1, total))
            )
            node = {
                "level": level,
                "title": hp["text"],
                "para_index": para_start,
                "children": children,
                "para_range": [para_start, para_end],
                "char_count": char_count,
            }
            return node, ci

        self._sections = []
        i = 0
        while i < len(heading_paras):
            hp = heading_paras[i]
            if hp["level"] == 1:
                node, i = build_node_single(i)
                self._sections.append(node)
            else:
                i += 1

    def _build_image_index(self):
        self._images = []
        image_rels = {}
        for rel in self.doc.part.rels.values():
            if "image" in rel.reltype:
                image_rels[rel.rId] = rel
        for info in self._para_index:
            for r_id in info.get("image_ids", []):
                if r_id in image_rels:
                    rel = image_rels[r_id]
                    img_info = {
                        "para_index": info["index"],
                        "r_id": r_id,
                        "filename": rel.target_ref,
                        "format": os.path.splitext(rel.target_ref)[1].lstrip("."),
                    }
                    try:
                        blob = rel.target_part.blob
                        img_info["size_bytes"] = len(blob)
                    except Exception:
                        img_info["size_bytes"] = None
                    self._fill_image_dimensions(img_info, info["index"])
                    self._fill_image_caption(img_info, info["index"])
                    self._images.append(img_info)

    def _fill_image_dimensions(self, img_info, para_index):
        para = self.doc.paragraphs[para_index]
        drawings = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        for drawing in drawings:
            extent = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
            if extent is None:
                continue
            cx = extent.get('cx')
            cy = extent.get('cy')
            if cx and cy:
                img_info["width_cm"] = emu_to_cm(int(cx))
                img_info["height_cm"] = emu_to_cm(int(cy))
                break
        docPr = para._element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr')
        if docPr is not None:
            img_info["name"] = docPr.get('name', '')
            img_info["description"] = docPr.get('descr', '')

    def _fill_image_caption(self, img_info, para_index):
        caption_pattern = re.compile(r'^\s*(图|表|Figure|Table)\s*\d+[.\-]\d+')
        search_range = 3
        best_match = None
        for offset in range(-search_range, search_range + 1):
            idx = para_index + offset
            if 0 <= idx < len(self._para_index):
                p = self._para_index[idx]
                text = p["text"].strip()
                if not text:
                    continue
                if p["style"].lower() == "caption":
                    img_info["nearby_caption"] = text
                    img_info["caption_para_index"] = idx
                    return
                if caption_pattern.match(text) and len(text) <= 60:
                    if best_match is None or len(text) > len(best_match[0]):
                        best_match = (text, idx)
        if best_match:
            img_info["nearby_caption"] = best_match[0]
            img_info["caption_para_index"] = best_match[1]
        else:
            img_info["nearby_caption"] = None
            img_info["caption_para_index"] = None

    def _build_table_index(self):
        self._tables = []
        for idx, table in enumerate(self.doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
            header = rows[0] if rows else []
            data = rows[1:] if len(rows) > 1 else []
            tbl_element = table._element
            para_index_approx = 0
            body = self.doc.element.body
            found = False
            for child_idx, child in enumerate(body):
                if child is tbl_element:
                    para_count = 0
                    for prev_child in body[:child_idx]:
                        if prev_child.tag.endswith('}p'):
                            para_count += 1
                    para_index_approx = para_count
                    found = True
                    break
            self._tables.append({
                "index": idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "header": header,
                "data": data,
                "para_index_approx": para_index_approx,
            })

    @property
    def paragraphs(self):
        return self._para_index

    @property
    def sections_tree(self):
        return self._sections

    @property
    def images(self):
        return self._images

    @property
    def tables(self):
        return self._tables

    @property
    def raw_paragraphs(self):
        return self.doc.paragraphs

    @property
    def body(self):
        """文档 body 元素（XML 操作入口）。"""
        return self.doc.element.body

    @property
    def raw_tables(self):
        return self.doc.tables

    # ── 修订操作 ──

    def accept_all_revisions(self):
        """接受全部修订标记（插入保留、删除移除）。"""
        from lib.reviser import accept_all_revisions as _accept
        _accept(self)
        self._build_index()

    def reject_all_revisions(self):
        """拒绝全部修订标记（插入移除、删除恢复）。"""
        from lib.reviser import reject_all_revisions as _reject
        _reject(self)
        self._build_index()

    def accept_revision(self, para_index, match_text):
        """接受一条修订（按段落索引 + 文本内容匹配）。"""
        from lib.reviser import accept_revision as _accept_one
        result = _accept_one(self, para_index, match_text)
        self._build_index()
        return result

    def reject_revision(self, para_index, match_text):
        """拒绝一条修订（按段落索引 + 文本内容匹配）。"""
        from lib.reviser import reject_revision as _reject_one
        result = _reject_one(self, para_index, match_text)
        self._build_index()
        return result

    # ── 段落操作 ──

    def set_paragraph_text(self, index, text):
        """清空段落现有所有 run，写入一段纯文本。保留段落样式。"""
        if index < 0 or index >= len(self.raw_paragraphs):
            raise IndexError(f"段落索引 {index} 超出范围 (0-{len(self.raw_paragraphs)-1})")
        from lxml import etree
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        w = lambda tag: f'{{{W_NS}}}{tag}'
        para = self.raw_paragraphs[index]
        p_elem = para._element
        for r in list(p_elem.findall(w('r'))):
            p_elem.remove(r)
        for ins in list(p_elem.findall(w('ins'))):
            p_elem.remove(ins)
        for d in list(p_elem.findall(w('del'))):
            p_elem.remove(d)
        run = etree.SubElement(p_elem, w('r'))
        t = etree.SubElement(run, w('t'))
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    def verify_index(self):
        """校验 _para_index 与 raw_paragraphs 长度一致。不一致时自动重建。"""
        if len(self._para_index) != len(self.raw_paragraphs):
            self._build_index()
            return False
        return True

    def find_section(self, title=None, level=None, index=None):
        def _search(nodes, results):
            for node in nodes:
                match = True
                if title and title not in node["title"]:
                    match = False
                if level is not None and node["level"] != level:
                    match = False
                if match:
                    results.append(node)
                _search(node["children"], results)
        results = []
        _search(self._sections, results)
        if index is not None and 1 <= index <= len(results):
            return results[index - 1]
        if title:
            return results[0] if results else None
        return results

    def get_section_paras(self, section_node):
        start, end = section_node["para_range"]
        return self._para_index[start:end + 1]

    def get_para(self, index):
        if 0 <= index < len(self._para_index):
            return self._para_index[index]
        return None

    def find_paragraph_by_text(self, text, start=0):
        for p in self._para_index[start:]:
            if text in p["text"]:
                return p["index"]
        return None

    def save(self, output_path=None):
        """保存文档。默认使用 save_zip（安全保留公式/图片），
        而非 python-docx 原生的 doc.save()（会丢失 OMML 公式和插入的图片）。"""
        return self.save_zip(output_path)

    # ── 上下文管理器 ──

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        # 不自动 save，调用方必须显式 save()。
        # 这样在 with 块内可多次操作，最后统一保存。
        pass

    def _build_rels_xml(self):
        """从内存中的 rels 构建 word/_rels/document.xml.rels XML。"""
        from lxml import etree
        NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
        root = etree.Element(f'{{{NS}}}Relationships')
        for rId, rel in self.doc.part.rels.items():
            child = etree.SubElement(root, f'{{{NS}}}Relationship')
            child.set('Id', rId)
            child.set('Type', rel.reltype)
            child.set('Target', rel.target_ref)
        return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

    def save_zip(self, output_path=None):
        import zipfile
        import tempfile
        from lxml import etree
        path = output_path or self.filepath
        xml_bytes = etree.tostring(
            self.doc.element,
            xml_declaration=True,
            encoding='UTF-8',
            standalone=True,
        )
        rels_xml = self._build_rels_xml()
        output_dir = os.path.dirname(os.path.abspath(path))
        tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx', dir=output_dir)
        os.close(tmp_fd)
        try:
            with zipfile.ZipFile(self.filepath, 'r') as zin:
                with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                    # 去重写入 zip 中的已有文件
                    written = set()
                    for item in zin.infolist():
                        if item.filename in written:
                            continue
                        written.add(item.filename)
                        if item.filename == 'word/document.xml':
                            zout.writestr(item, xml_bytes)
                        elif item.filename == 'word/_rels/document.xml.rels':
                            zout.writestr(item, rels_xml)
                        else:
                            zout.writestr(item, zin.read(item))
                    # 写入新 blob（insert_image / replace_image 添加的内存中图片）
                    # 以及内存中修改过的 header/footer 等 XML part
                    for rel in self.doc.part.rels.values():
                        if not hasattr(rel, 'target_part'):
                            continue
                        target_ref = rel.target_ref
                        if target_ref.startswith('/'):
                            target_ref = target_ref[1:]
                        elif target_ref.startswith('../'):
                            target_ref = target_ref[3:]
                        # 统一补上 word/ 前缀（zip 内文件的实际路径）
                        if not target_ref.startswith('word/'):
                            target_ref = f'word/{target_ref}'
                        if target_ref in written:
                            continue
                        written.add(target_ref)
                        try:
                            if hasattr(rel.target_part, 'blob'):
                                zout.writestr(target_ref, rel.target_part.blob)
                            elif hasattr(rel.target_part, '_element'):
                                part_xml = etree.tostring(
                                    rel.target_part._element,
                                    xml_declaration=True, encoding='UTF-8', standalone=True)
                                zout.writestr(target_ref, part_xml)
                        except Exception as e:
                            import warnings
                            warnings.warn(f'save_zip: 无法写入 {target_ref}: {e}')
            os.replace(tmp_path, path)
        except Exception:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
            raise
        return path
