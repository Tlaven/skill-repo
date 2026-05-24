"""编辑模块 — 纯库函数，所有函数接受显式参数而非 argparse.Namespace"""
import json
import os
import zipfile
import tempfile
from lxml import etree
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lib.utils import NSMAP, get_output_path
from lib.styles import STYLE_NAME_TO_WORD
from lib.fixer import clear_direct_formatting, ensure_word_styles


def _clear_para_runs(para):
    """移除段落的所有 run、修订插入（ins）和修订删除（del）元素。"""
    from lxml import etree as _et
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p_elem = para._element
    for tag in (f'{{{W}}}r', f'{{{W}}}ins', f'{{{W}}}del'):
        for child in list(p_elem.findall(tag)):
            p_elem.remove(child)


def replace_text(doc, paragraph, text, output=None, backup=False):
    """替换段落文字。保持段落样式不变。"""
    output_path = get_output_path(doc, output=output, backup=backup)
    para = doc.raw_paragraphs[paragraph]
    old_text = para.text or ""
    _clear_para_runs(para)
    para.add_run(text)
    doc.save_zip(output_path)
    return {"paragraph": paragraph, "old_text": old_text, "new_text": text, "output": output_path}


def replace_batch(doc, pairs, chapter=None, output=None, backup=False):
    """批量替换。pairs: [{"old": "旧词", "new": "新词"}, ...]"""
    output_path = get_output_path(doc, output=output, backup=backup)
    para_range = None
    if chapter is not None:
        section = doc.find_section(level=1, index=chapter)
        if section is None:
            return {"error": f"未找到第 {chapter} 章"}
        para_range = section["para_range"]
    total_replacements = 0
    details = []
    for pair in pairs:
        old = pair.get("old", ""); new = pair.get("new", "")
        count = 0
        for i, para_info in enumerate(doc.paragraphs):
            if para_range and (i < para_range[0] or i > para_range[1]):
                continue
            para = doc.raw_paragraphs[i]
            if old not in para.text:
                continue
            replaced = _replace_in_paragraph(para, old, new)
            if replaced:
                count += 1; total_replacements += 1
        details.append({"old": old, "new": new, "replacements": count})
    doc.save_zip(output_path)
    return {"total_replacements": total_replacements, "details": details, "output": output_path}


def replace_batch_by_index(doc, pairs_file, output=None, backup=False):
    """按段落索引批量替换。pairs_file: {"43": "新文本", ...} JSON 文件。"""
    output_path = get_output_path(doc, output=output, backup=backup)
    with open(pairs_file, 'r', encoding='utf-8') as f:
        pairs = json.load(f)
    if not isinstance(pairs, dict) or not pairs:
        return {"error": "pairs-file 格式错误：需要 JSON 对象"}
    details = []
    for idx_str, new_text in sorted(pairs.items(), key=lambda x: int(x[0])):
        idx = int(idx_str)
        if idx < 0 or idx >= len(doc.raw_paragraphs):
            details.append({"paragraph": idx, "error": f"索引超出范围 (0-{len(doc.raw_paragraphs)-1})"})
            continue
        para = doc.raw_paragraphs[idx]
        old_text = para.text or ""
        _clear_para_runs(para)
        para.add_run(new_text)
        details.append({"paragraph": idx, "old_chars": len(old_text), "new_chars": len(new_text)})
    doc.save_zip(output_path)
    return {"total_replaced": len([d for d in details if "error" not in d]), "details": details, "output": output_path}


def _replace_in_paragraph(para, old_text, new_text):
    """在段落中进行文本替换。"""
    if not para.runs or old_text not in para.text:
        return False
    full_text = para.text
    if old_text not in full_text:
        return False
    new_full = full_text.replace(old_text, new_text)
    _clear_para_runs(para)
    p_element = para._element
    new_run_elem = etree.SubElement(p_element, f'{{{NSMAP["w"]}}}r')
    t = etree.SubElement(new_run_elem, f'{{{NSMAP["w"]}}}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = new_full
    return True


def replace_inline(doc, paragraph, old, new, output=None, backup=False,
                   bold=None, font=None, font_east=None, size=None, color=None):
    """段内子串替换，不替换全段。

    支持 --deep 格式参数：bold/font/font_east/size/color。
    指定格式参数时，将应用到被替换文本所在的 run。
    """
    output_path = get_output_path(doc, output=output, backup=backup)
    para = doc.raw_paragraphs[paragraph]
    full_text = para.text or ""
    if old not in full_text:
        return {"paragraph": paragraph, "error": f"未找到子串 '{old[:30]}'"}
    count = full_text.count(old)
    has_format = any(x is not None for x in (bold, font, font_east, size, color))

    pos = 0
    replaced = False
    for run in para.runs:
        if not run.text:
            continue
        if not replaced and old in run.text:
            run.text = run.text.replace(old, new, 1)
            if has_format:
                _apply_run_format(run, bold=bold, font=font, font_east=font_east,
                                  size=size, color=color)
            replaced = True
            break
        pos += len(run.text)
    if not replaced:
        doc_para = doc.paragraphs[paragraph]
        old_full = doc_para["text"]
        _clear_para_runs(para)
        para.add_run(old_full.replace(old, new, 1))
    doc.save_zip(output_path)
    return {"paragraph": paragraph, "old_substr": old, "new_substr": new,
            "replacements": count, "formatted": has_format}


def format_inline(doc, paragraph, target, output=None, backup=False,
                  bold=None, font=None, font_east=None, size=None, color=None):
    """不改文字，仅修改段内子串的格式（加粗/字体/字号/颜色）。"""
    output_path = get_output_path(doc, output=output, backup=backup)
    para = doc.raw_paragraphs[paragraph]
    full_text = para.text or ""
    if target not in full_text:
        return {"paragraph": paragraph, "error": f"未找到子串 '{target[:30]}'"}

    if not any(x is not None for x in (bold, font, font_east, size, color)):
        return {"paragraph": paragraph, "error": "请指定至少一种格式: --bold/--font/--size/--color"}

    pos = 0
    formatted = False
    for run in para.runs:
        if not run.text:
            continue
        if not formatted and target in run.text:
            _apply_run_format(run, bold=bold, font=font, font_east=font_east,
                              size=size, color=color)
            formatted = True
            break
        pos += len(run.text)
    if not formatted:
        return {"paragraph": paragraph, "error": f"在 runs 中未找到子串 '{target[:30]}'"}

    doc.save_zip(output_path)
    return {"paragraph": paragraph, "target": target, "formats_applied": {
        "bold": bold, "font": font, "font_east": font_east,
        "size": size, "color": color,
    }}


def _apply_run_format(run, bold=None, font=None, font_east=None, size=None, color=None):
    """设置 run 的格式属性。None 表示不修改。"""
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        from docx.shared import RGBColor
        try:
            run.font.color.rgb = RGBColor.from_string(color.lstrip('#'))
        except Exception:
            pass
    if size is not None:
        from docx.shared import Pt
        run.font.size = Pt(float(size))
    # 字体需通过 XML 设置东亚字体
    if font is not None or font_east is not None:
        rPr = run._element.find(f'{{{NSMAP["w"]}}}rPr')
        if rPr is None:
            from lxml import etree as _et
            rPr = _et.SubElement(run._element, f'{{{NSMAP["w"]}}}rPr')
        rFonts = rPr.find(f'{{{NSMAP["w"]}}}rFonts')
        if rFonts is None:
            from lxml import etree as _et
            rFonts = _et.SubElement(rPr, f'{{{NSMAP["w"]}}}rFonts')
        if font:
            rFonts.set(f'{{{NSMAP["w"]}}}ascii', font)
            rFonts.set(f'{{{NSMAP["w"]}}}hAnsi', font)
        if font_east:
            rFonts.set(f'{{{NSMAP["w"]}}}eastAsia', font_east)


def insert_paragraph(doc, after, text, style='body', rules=None, output=None, backup=False):
    """在指定位置插入段落。"""
    output_path = get_output_path(doc, output=output, backup=backup)
    ref_para = doc.raw_paragraphs[after]
    word_style = STYLE_NAME_TO_WORD.get(style, style)
    ensure_word_styles(doc.doc, {word_style}, rules)
    new_para = _create_clean_paragraph(text, word_style)
    ref_para._element.addnext(new_para)
    doc.save_zip(output_path)
    doc._build_index()
    return {"after_paragraph": after, "text": text, "style": word_style, "output": output_path}


def write_paragraphs(doc, after, data, output=None, backup=False):
    """批量写入多段（从后往前插入，防索引漂移）。

    data: [{"text": "...", "style": "body"}, ...]
    """
    output_path = get_output_path(doc, output=output, backup=backup)
    styles_needed = set()
    for item in data:
        s = item.get("style", "body")
        styles_needed.add(STYLE_NAME_TO_WORD.get(s, s))
    ensure_word_styles(doc.doc, styles_needed, None)

    inserted = []
    for item in reversed(data):
        text = item.get("text", "")
        style = item.get("style", "body")
        word_style = STYLE_NAME_TO_WORD.get(style, style)
        new_para = _create_clean_paragraph(text, word_style)
        ref_para = doc.raw_paragraphs[after]
        ref_para._element.addnext(new_para)
        inserted.insert(0, {"text": text[:50], "style": word_style})
    doc.save_zip(output_path)
    doc._build_index()
    return {"after_paragraph": after, "total_inserted": len(data), "inserted": inserted, "output": output_path}


def _create_clean_paragraph(text, word_style_name):
    from docx.oxml import parse_xml
    W = NSMAP["w"]
    xml_space = '{http://www.w3.org/XML/1998/namespace}space'
    p_raw = etree.Element(f'{{{W}}}p')
    pPr = etree.SubElement(p_raw, f'{{{W}}}pPr')
    pStyle = etree.SubElement(pPr, f'{{{W}}}pStyle')
    pStyle.set(f'{{{W}}}val', word_style_name)
    r = etree.SubElement(p_raw, f'{{{W}}}r')
    t = etree.SubElement(r, f'{{{W}}}t')
    t.set(xml_space, 'preserve')
    t.text = text
    p_xml = etree.tostring(p_raw, encoding='unicode')
    return parse_xml(p_xml)


def delete_paragraph(doc, paragraph, output=None, backup=False):
    output_path = get_output_path(doc, output=output, backup=backup)
    para = doc.raw_paragraphs[paragraph]
    deleted_text = para.text or ""
    para._element.getparent().remove(para._element)
    doc.save_zip(output_path)
    doc._build_index()
    return {"deleted_paragraph": paragraph, "deleted_text": deleted_text, "output": output_path}


def set_format(doc, style, paragraph=None, start=None, end=None, target=None, rules=None, output=None, backup=False):
    output_path = get_output_path(doc, output=output, backup=backup)
    if not style:
        return {"error": "请指定 --style <name>"}
    word_style = STYLE_NAME_TO_WORD.get(style, style)
    ensure_word_styles(doc.doc, {word_style}, rules)
    target_indices = _resolve_format_targets(doc, paragraph=paragraph, start=start, end=end, target=target)
    if isinstance(target_indices, dict) and "error" in target_indices:
        return target_indices
    modified = 0
    for para_idx in target_indices:
        para = doc.raw_paragraphs[para_idx]
        try:
            para.style = doc.doc.styles[word_style]
        except KeyError:
            continue
        clear_direct_formatting(para)
        modified += 1
    doc.save_zip(output_path)
    return {"style": word_style, "total_paragraphs": len(target_indices), "modified": modified, "output": output_path}


def _resolve_format_targets(doc, paragraph=None, start=None, end=None, target=None):
    if paragraph is not None:
        return [paragraph]
    if start is not None and end is not None:
        return list(range(start, end + 1))
    if target:
        from lib.styles import classify_paragraph
        indices = []
        for p in doc.paragraphs:
            text = p["text"].strip()
            if not text: continue
            if target == 'body':
                if p["style"] not in ("Normal", "Body Text"): continue
                if classify_paragraph(text) is not None: continue
                indices.append(p["index"])
            elif target == 'headings':
                if p["level"] is not None:
                    indices.append(p["index"])
        if not indices:
            return {"error": f"未找到匹配 --target {target} 的段落"}
        return indices
    return {"error": "请指定 --paragraph N 或 --start N --end M 或 --target body|headings"}


def replace_table(doc, index, data, output=None, backup=False):
    output_path = get_output_path(doc, output=output, backup=backup)
    if index < 0 or index >= len(doc.raw_tables):
        return {"error": f"表格索引 {index} 超出范围 (0-{len(doc.raw_tables)-1})"}
    if not data or not isinstance(data[0], list):
        return {"error": "数据格式错误：需要二维数组"}
    table = doc.raw_tables[index]
    num_rows = len(data); num_cols = max(len(row) for row in data)
    current_rows = len(table.rows)
    if current_rows < num_rows:
        for _ in range(num_rows - current_rows):
            table.add_row()
    elif current_rows > num_rows:
        tbl_element = table._tbl
        tr_elements = tbl_element.findall(f'{{{NSMAP["w"]}}}tr')
        for tr in tr_elements[num_rows:]:
            tbl_element.remove(tr)
    current_cols = len(table.columns)
    if current_cols < num_cols:
        for _ in range(num_cols - current_cols):
            table.add_column()
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.cell(row_idx, col_idx)
                for cell_para in cell.paragraphs:
                    _clear_para_runs(cell_para)
                    cell_para._element.clear()
                    new_run = etree.SubElement(paragraph._element, f'{{{NSMAP["w"]}}}r')
                    t = etree.SubElement(new_run, f'{{{NSMAP["w"]}}}t')
                    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    t.text = str(cell_text) if cell_text is not None else ""
    doc.save_zip(output_path)
    return {"table_index": index, "rows": num_rows, "cols": num_cols,
            "data_preview": [row[:3] for row in data[:3]], "output": output_path}


def insert_table(doc, after, data, output=None, backup=False):
    output_path = get_output_path(doc, output=output, backup=backup)
    if not data or not isinstance(data[0], list):
        return {"error": "数据格式错误：需要二维数组"}
    num_rows = len(data); num_cols = max(len(row) for row in data)
    table = doc.doc.add_table(rows=num_rows, cols=num_cols)
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                table.cell(row_idx, col_idx).text = str(cell_text) if cell_text is not None else ""
    ref_para = doc.raw_paragraphs[after]
    tbl_element = table._tbl
    tbl_element.getparent().remove(tbl_element)
    ref_para._element.addnext(tbl_element)
    doc.save_zip(output_path)
    doc._build_index()
    return {"after_paragraph": after, "rows": num_rows, "cols": num_cols,
            "data_preview": [row[:3] for row in data[:3]], "output": output_path}


def _get_caption_style_id(doc):
    """从 styles.xml 查找实际定义的 caption/题注 样式 styleId。"""
    from lxml import etree as _etree
    W = NSMAP['w']
    try:
        import zipfile
        with zipfile.ZipFile(doc.filepath, 'r') as z:
            if 'word/styles.xml' not in z.namelist():
                return None
            styles_xml = _etree.fromstring(z.read('word/styles.xml'))
            for style in styles_xml.iter(f'{{{W}}}style'):
                sid = style.get(f'{{{W}}}styleId')
                name_el = style.find(f'{{{W}}}name')
                nval = name_el.get(f'{{{W}}}val') if name_el is not None else ''
                if nval in ('Caption', 'caption', '题注'):
                    return sid
    except Exception:
        pass
    return None


def _detect_caption_style(doc):
    """检测文档中图题/表题实际使用的样式 styleId。先查 styles.xml 定义，再扫正文段落。"""
    # 优先：从 styles.xml 获取真实的 caption 样式 ID
    defined_id = _get_caption_style_id(doc)
    # 如果定义了且不是 Normal，直接返回
    if defined_id and defined_id not in ('Normal', ''):
        return defined_id
    # 备选：扫已有段落，匹配实际使用的样式
    from lxml import etree as _etree
    W = NSMAP['w']
    body = doc.doc.element.find(f'{{{W}}}body')
    if body is not None:
        for p in body.iter(f'{{{W}}}p'):
            texts = p.findall(f'.//{{{W}}}t')
            full = ''.join(t.text or '' for t in texts)
            stripped = full.strip()
            if not (stripped.startswith('图') or stripped.startswith('表')):
                continue
            pPr = p.find(f'{{{W}}}pPr')
            if pPr is not None:
                pStyle = pPr.find(f'{{{W}}}pStyle')
                if pStyle is not None:
                    sid = pStyle.get(f'{{{W}}}val')
                    if sid and sid not in ('Normal', ''):
                        return sid
    return 'Caption'


def insert_image(doc, after, image, width=None, caption=None, output=None, backup=False):
    """在指定段落后插入图片。使用 save_zip 保存以保留已有 OMML 公式。"""
    output_path = get_output_path(doc, output=output, backup=backup)
    if not os.path.exists(image):
        return {"error": f"图片文件不存在: {image}"}
    ext = os.path.splitext(image)[1].lower()
    if ext == '.svg':
        return {"error": f"不支持 SVG 格式图片（python-docx 限制）。请将图片转换为 PNG/JPEG 后再插入。"}
    section = doc.doc.sections[0]
    page_width_cm = section.page_width / 360000
    margin_left_cm = section.left_margin / 360000
    margin_right_cm = section.right_margin / 360000
    margin_top_cm = section.top_margin / 360000
    margin_bottom_cm = section.bottom_margin / 360000
    page_height_cm = section.page_height / 360000
    text_width_cm = page_width_cm - margin_left_cm - margin_right_cm
    max_height_cm = page_height_cm - margin_top_cm - margin_bottom_cm - 2
    default_width_cm = text_width_cm * 0.8
    if width:
        try:
            w = float(width)
            if w > text_width_cm: w = text_width_cm
            width_cm = Cm(w)
        except ValueError:
            return {"error": f"无效的宽度值: {width}"}
    else:
        width_cm = Cm(default_width_cm)
    p = doc.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    inline_shape = run.add_picture(image, width=width_cm)
    actual_height_cm = inline_shape.height / 360000
    if actual_height_cm > max_height_cm:
        ratio = max_height_cm / actual_height_cm
        new_width_cm = (inline_shape.width / 360000) * ratio
        inline_shape.width = Cm(new_width_cm)
        inline_shape.height = Cm(max_height_cm)
    actual_width_cm = inline_shape.width / 360000
    actual_height_cm = inline_shape.height / 360000
    target_para = doc.raw_paragraphs[after]
    target_para._element.addnext(p._element)
    if caption:
        cap_style = _detect_caption_style(doc)
        cap_p = _create_clean_paragraph(caption, cap_style)
        p._element.addnext(cap_p)
    doc.save_zip(output_path)
    doc._build_index()
    result = {
        "after_paragraph": after, "image": os.path.basename(image),
        "output": output_path, "width_cm": f"{actual_width_cm:.1f}",
        "height_cm": f"{actual_height_cm:.1f}",
    }
    if caption:
        result["caption"] = caption
    return result


def replace_image(doc, image, caption=None, paragraph=None, media=None, output=None, backup=False):
    """替换文档中已有的图片。使用 save_zip 保存以保留 OMML 公式。"""
    output_path = get_output_path(doc, output=output, backup=backup)
    if not os.path.exists(image):
        return {"error": f"图片文件不存在: {image}"}
    target = _resolve_image_target(doc, caption=caption, paragraph=paragraph, media=media)
    if target is None:
        return {"error": "未找到匹配的图片。请用 --caption / --paragraph / --media 指定目标。"}
    media_entry, matched_info = target
    with open(image, 'rb') as f:
        new_blob = f.read()
    from lxml import etree as _etree
    xml_bytes = _etree.tostring(doc.doc.element, xml_declaration=True, encoding='UTF-8', standalone=True)
    output_dir = os.path.dirname(os.path.abspath(output_path))
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx', dir=output_dir)
    os.close(tmp_fd)
    try:
        with zipfile.ZipFile(doc.filepath, 'r') as zin:
            with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.namelist():
                    if item == 'word/document.xml':
                        zout.writestr(item, xml_bytes)
                    elif item == media_entry:
                        zout.writestr(item, new_blob)
                    else:
                        zout.writestr(item, zin.read(item))
        os.replace(tmp_path, output_path)
    except Exception:
        if os.path.exists(tmp_path): os.unlink(tmp_path)
        raise
    return {
        "replaced": media_entry, "new_image": os.path.basename(image),
        "old_size": matched_info.get("size_bytes"), "new_size": len(new_blob),
        "matched_by": matched_info["method"], "output": output_path,
    }


def _resolve_image_target(doc, caption=None, paragraph=None, media=None):
    if caption:
        normalized = caption.replace(' ', '').replace('\u00a0', '')
        for img in doc.images:
            nc = img.get("nearby_caption")
            if nc:
                nc_norm = nc.replace(' ', '').replace('\u00a0', '')
                if normalized in nc_norm or nc_norm in normalized:
                    rid = img["r_id"]
                entry = _rid_to_media_entry(doc, rid)
                if entry:
                    return entry, {"method": f"caption '{caption}'", "size_bytes": img.get("size_bytes")}
        return None
    if paragraph is not None:
        if paragraph < 0 or paragraph >= len(doc.paragraphs):
            return None
        pi = doc.paragraphs[paragraph]
        if not pi["has_image"] or not pi["image_ids"]:
            return None
        rid = pi["image_ids"][0]
        entry = _rid_to_media_entry(doc, rid)
        if entry:
            return entry, {"method": f"paragraph {paragraph}", "size_bytes": _get_image_size(doc, rid)}
        return None
    if media:
        for img in doc.images:
            if media in img["filename"] or img["filename"].endswith(media):
                rid = img["r_id"]
                entry = _rid_to_media_entry(doc, rid)
                if entry:
                    return entry, {"method": f"media '{media}'", "size_bytes": img.get("size_bytes")}
        return None
    return None


def _rid_to_media_entry(doc, rid):
    try:
        rel = doc.doc.part.rels[rid]
        ref = rel.target_ref
        if ref.startswith('/'):
            return ref[1:]
        if '/' not in ref:
            return f"word/media/{ref}"
        return ref
    except KeyError:
        return None


def _get_image_size(doc, rid):
    try:
        rel = doc.doc.part.rels[rid]
        return len(rel.target_part.blob)
    except Exception:
        return None



