"""页面布局操作模块 — 纯库函数"""
import re
from lxml import etree
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lib.utils import NSMAP, get_output_path
from lib.styles import resolve_style


def set_page_setup(doc, width=None, height=None, margin_top=None, margin_bottom=None, margin_left=None, margin_right=None, output=None, backup=False):
    output_path = get_output_path(doc, output=output, backup=backup)
    changes = {}
    for section in doc.doc.sections:
        if width: section.page_width = Cm(float(width)); changes["width_cm"] = float(width)
        if height: section.page_height = Cm(float(height)); changes["height_cm"] = float(height)
        if margin_top: section.top_margin = Cm(float(margin_top)); changes["margin_top_cm"] = float(margin_top)
        if margin_bottom: section.bottom_margin = Cm(float(margin_bottom)); changes["margin_bottom_cm"] = float(margin_bottom)
        if margin_left: section.left_margin = Cm(float(margin_left)); changes["margin_left_cm"] = float(margin_left)
        if margin_right: section.right_margin = Cm(float(margin_right)); changes["margin_right_cm"] = float(margin_right)
    doc.save(output_path)
    return {"changes": changes, "sections_modified": len(doc.doc.sections), "output": output_path}


def insert_page_break(doc, after, output=None, backup=False):
    output_path = get_output_path(doc, output=output, backup=backup)
    ref_para = doc.raw_paragraphs[after]
    new_p = etree.SubElement(ref_para._element.getparent(), f'{{{NSMAP["w"]}}}p')
    new_p.append(etree.SubElement(new_p, f'{{{NSMAP["w"]}}}r'))
    br = etree.SubElement(new_p.find(f'{{{NSMAP["w"]}}}r'), f'{{{NSMAP["w"]}}}br')
    br.set(f'{{{NSMAP["w"]}}}type', 'page')
    ref_para._element.addnext(new_p)
    doc.save(output_path)
    return {"after_paragraph": after, "output": output_path,
            "note": "段落索引已偏移，后续操作前请先 read-structure 获取新索引"}


def set_header(doc, text, font='宋体', size='9', output=None, backup=False):
    output_path = get_output_path(doc, output=output, backup=backup)
    header_style = resolve_style("header")
    font_name = font or header_style.get("font_east", "宋体")
    font_size = float(size or header_style.get("size_pt", 9))
    for section in doc.doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        for para in header.paragraphs:
            for run in para.runs:
                run._element.getparent().remove(run._element)
            para._element.clear()
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        align = header_style.get("alignment", "center")
        para.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        rPr = run._element.find(f'{{{NSMAP["w"]}}}rPr')
        if rPr is not None:
            rFonts = rPr.find(f'{{{NSMAP["w"]}}}rFonts')
            if rFonts is not None:
                rFonts.set(f'{{{NSMAP["w"]}}}eastAsia', font_name)
    doc.save(output_path)
    return {"header_text": text, "font": font_name, "size_pt": font_size,
            "sections_modified": len(doc.doc.sections), "output": output_path}


def set_footer(doc, text=None, page_number=False, align='center', font='宋体', size='9', output=None, backup=False):
    output_path = get_output_path(doc, output=output, backup=backup)
    footer_style = resolve_style("footer")
    align = align or footer_style.get("alignment", "center")
    font_name = font or footer_style.get("font_east", "宋体")
    font_size = float(size or footer_style.get("size_pt", 9))
    for section in doc.doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        for para in footer.paragraphs:
            for run in para.runs:
                run._element.getparent().remove(run._element)
            para._element.clear()
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
        if page_number:
            run = para.add_run()
            run.font.name = font_name; run.font.size = Pt(font_size)
            fld_char_begin = etree.SubElement(run._element, f'{{{NSMAP["w"]}}}fldChar')
            fld_char_begin.set(f'{{{NSMAP["w"]}}}fldCharType', 'begin')
            instr_run = para.add_run()
            instr_run.font.name = font_name; instr_run.font.size = Pt(font_size)
            instr_text = etree.SubElement(instr_run._element, f'{{{NSMAP["w"]}}}instrText')
            instr_text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            instr_text.text = ' PAGE '
            end_run = para.add_run()
            end_run.font.name = font_name; end_run.font.size = Pt(font_size)
            fld_char_end = etree.SubElement(end_run._element, f'{{{NSMAP["w"]}}}fldChar')
            fld_char_end.set(f'{{{NSMAP["w"]}}}fldCharType', 'end')
        elif text:
            run = para.add_run(text)
            run.font.name = font_name; run.font.size = Pt(font_size)
    doc.save(output_path)
    return {"page_number": page_number, "footer_text": text, "align": align, "output": output_path}


def _build_chapter_map(doc):
    """构建 段落索引 → 实际章号 的映射（从标题文字"第X章"解析，而非内部索引位置）"""
    chapter_map = {}
    current_chapter = 0
    ch_pattern = re.compile(r'第(\d+)章')
    for p in doc.paragraphs:
        if p.get("level") == 1:
            text = p.get("text", "")
            m = ch_pattern.search(text)
            if m:
                current_chapter = int(m.group(1))
            else:
                m2 = re.match(r'(\d+)[.、．\s]', text)
                if m2:
                    current_chapter = int(m2.group(1))
        chapter_map[p["index"]] = current_chapter
    return chapter_map


def renumber_figures(doc, output=None, backup=False):
    output_path = get_output_path(doc, output=output, backup=backup)
    fig_pattern = re.compile(r'(图\s*)(\d+)([-.]\s*)(\d+)')
    tbl_pattern = re.compile(r'(表\s*)(\d+)([-.]\s*)(\d+)')
    chapter_map = _build_chapter_map(doc)
    chapter_fig_counters = {}
    chapter_tbl_counters = {}
    changes = []
    for p_info in doc.paragraphs:
        text = p_info.get("text", "")
        para_idx = p_info["index"]
        para = doc.raw_paragraphs[para_idx]
        for pattern, counters, label in [
            (fig_pattern, chapter_fig_counters, "图"),
            (tbl_pattern, chapter_tbl_counters, "表"),
        ]:
            for match in pattern.finditer(text):
                old_chapter = int(match.group(2))
                old_num = int(match.group(4))
                current_chapter = chapter_map.get(para_idx, old_chapter)
                counters.setdefault(current_chapter, 0)
                counters[current_chapter] += 1
                new_num = counters[current_chapter]
                if current_chapter != old_chapter or new_num != old_num:
                    old_label = match.group(0)
                    new_label = f"{match.group(1)}{current_chapter}{match.group(3)}{new_num}"
                    _replace_in_runs(para, old_label, new_label)
                    changes.append({"para_index": para_idx, "old": old_label, "new": new_label})
    if changes:
        doc.save(output_path)
    return {
        "total_renumbered": len(changes), "changes": changes[:50],
        "figure_counters": chapter_fig_counters, "table_counters": chapter_tbl_counters,
        "output": output_path if changes else None,
    }


def _replace_in_runs(para, old_text, new_text):
    for run in para.runs:
        if run.text and old_text in run.text:
            run.text = run.text.replace(old_text, new_text, 1)
            return True
    return False
