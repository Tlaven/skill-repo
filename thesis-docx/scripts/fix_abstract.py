# -*- coding: utf-8 -*-
"""Fix abstract section: accept revisions, rewrite abstract, sync English abstract.
Usage: python scripts/fix_abstract.py <path_to_docx>
"""
import sys, os, shutil, copy
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from lxml import etree
from docx import Document
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from lib.core import ThesisDoc

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
w = lambda tag: f'{{{W_NS}}}{tag}'

def accept_all_revisions(doc):
    """Accept all tracked changes in the document."""
    body = doc.element.body
    for p_elem in body.iter(w('p')):
        for ins in list(p_elem.findall(w('ins'))):
            parent = ins.getparent()
            idx = list(parent).index(ins)
            for child in list(ins):
                parent.insert(idx, child)
                idx += 1
            parent.remove(ins)
        for d in list(p_elem.findall(w('del'))):
            d.getparent().remove(d)
    for tbl in body.iter(w('tbl')):
        for p_elem in tbl.iter(w('p')):
            for ins in list(p_elem.findall(w('ins'))):
                parent = ins.getparent()
                idx = list(parent).index(ins)
                for child in list(ins):
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(ins)
            for d in list(p_elem.findall(w('del'))):
                d.getparent().remove(d)

def get_para_text(para):
    return ''.join(t.text or '' for t in para._element.iter(w('t')))

def set_para_text(para, text):
    """Replace all content of a paragraph with a single run of text."""
    elem = para._element
    for r in list(elem.findall(w('r'))):
        elem.remove(r)
    for ins in list(elem.findall(w('ins'))):
        elem.remove(ins)
    for d in list(elem.findall(w('del'))):
        elem.remove(d)
    run = etree.SubElement(elem, w('r'))
    t = etree.SubElement(run, w('t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

def delete_para(para):
    para._element.getparent().remove(para._element)

NEW_CN = (
    "针对军队人才管理中多源异构数据利用不足、人岗匹配效率低的问题，"
    "提出一种基于大语言模型的人才画像与智能匹配方法。"
    "首先，设计面向警务（军队）人员的数据模型与多源异构数据预处理流程，"
    "实现35个字段的结构化提取与标准化。"
    "其次，采用PCA降维（保留90%方差贡献率）结合K-Means++聚类将人员自动划分为5类群体，"
    "利用大语言模型从群体到个体逐层生成结构化画像并进行向量化存储。"
    "最后，设计「任务语义解析—群体级余弦粗筛—个体级LLM深度评估」三级匹配策略，"
    "实现端到端智能匹配。"
    "在5000条警务代理数据集上的实验表明：聚类轮廓系数达0.52（K=5），"
    "画像标签匹配准确率超85%，个体评分与履历相关系数r=0.73；"
    "LLM智能匹配模式Top-5准确率76.7%，相比规则模式（53.3%）提升23.4个百分点，"
    "MRR和NDCG分别达0.81和0.84。"
    "受限于军队数据涉密性，实验采用警务代理数据集，后续需在真实数据上进一步验证。"
)

NEW_EN = (
    "This paper addresses the issues of insufficient utilization of multi-source heterogeneous data "
    "and low efficiency in person-job matching for military talent management by proposing a talent "
    "profiling and intelligent matching method based on large language models. "
    "First, a data model and preprocessing pipeline tailored for police (military) personnel is designed, "
    "enabling structured extraction and standardization of 35 fields. "
    "Second, PCA (retaining 90% variance contribution) combined with K-Means++ clustering automatically "
    "partitions personnel into 5 groups, and large language models generate structured profiles from "
    "group to individual levels, which are vectorized and stored. "
    "Finally, a three-tier matching strategy—task semantic parsing, group-level cosine粗筛, and "
    "individual-level LLM deep evaluation—achieves end-to-end intelligent matching. "
    "Experiments on a 5,000-sample police proxy dataset show: a silhouette coefficient of 0.52 (K=5), "
    "profile tag matching accuracy exceeding 85%, an average correlation of r=0.73 between individual "
    "scores and r\u00e9sum\u00e9s. The LLM-based matching mode achieves a Top-5 accuracy of 76.7%, a 23.4 percentage "
    "point improvement over the rule-based mode (53.3%), with MRR and NDCG reaching 0.81 and 0.84, respectively."
)

def find_para_by_text(paras, keyword):
    for i, p in enumerate(paras):
        if keyword in get_para_text(p):
            return i
    return None

def main(filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        sys.exit(1)
    filepath = os.path.abspath(filepath)
    backup = filepath.replace('.docx', f'_backup.docx')
    shutil.copy2(filepath, backup)
    print(f"Backup: {backup}")

    # Step 1: Load with python-docx, accept all revisions
    doc = Document(filepath)
    accept_all_revisions(doc)

    # Step 2: Save to temp file and reload to get fresh indices
    tmp = filepath.replace('.docx', '_tmp.docx')
    doc.save(tmp)
    doc2 = Document(tmp)

    paras = doc2.paragraphs

    cn_idx = find_para_by_text(paras, "本文针对军队人才管理")
    en_idx = find_para_by_text(paras, "This paper addresses")
    heading_idx = find_para_by_text(paras, "摘")
    template_label_idx = find_para_by_text(paras, "摘要参考模板")
    template_idx = find_para_by_text(paras, "医疗数据碎片化")

    print(f"CN abstract: para {cn_idx}")
    print(f"EN abstract: para {en_idx}")
    print(f"Heading: para {heading_idx}")
    print(f"Template label: para {template_label_idx}")
    print(f"Template: para {template_idx}")

    # Step 3: Replace Chinese abstract
    if cn_idx is not None:
        set_para_text(paras[cn_idx], NEW_CN)
        print("CN abstract replaced.")

    # Step 4: Replace English abstract
    if en_idx is not None:
        set_para_text(paras[en_idx], NEW_EN)
        print("EN abstract replaced.")

    # Step 5: Clean up template artifacts
    # Delete from back to front to preserve indices
    to_delete = []
    if template_idx is not None:
        to_delete.append(template_idx)
    if template_label_idx is not None:
        to_delete.append(template_label_idx)

    # Remove empty paragraphs between heading and keywords
    found_heading = False
    for i, p in enumerate(paras):
        text = get_para_text(p).strip()
        style = p.style.name if p.style else ""
        if "摘" in text and style == "Heading 1":
            found_heading = True
            continue
        if "关键词" in text or "Keywords" in text:
            break
        if found_heading and text == "":
            to_delete.append(i)

    for i in sorted(set(to_delete), reverse=True):
        delete_para(paras[i])
        print(f"Deleted para {i}")

    # Step 6: Save using save_zip
    doc2.save(tmp)
    # Reload with ThesisDoc for safe zip save
    thesis = ThesisDoc(tmp)
    thesis.save_zip(filepath)
    print(f"Saved to {filepath}")

    # Cleanup temp
    for f in [tmp]:
        if os.path.exists(f):
            os.remove(f)

    print("Done!")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python scripts/fix_abstract.py <path_to_docx>")
        sys.exit(1)
    main(sys.argv[1])
