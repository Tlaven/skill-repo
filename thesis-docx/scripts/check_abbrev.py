"""
扫描 docx 中英文简称首次出现是否遵循「中文（English Full Name, ABBR）」规范。

用法:
    python check_abbrev.py "论文.docx"
    python check_abbrev.py "论文.docx" -o report.json
    python check_abbrev.py "论文.docx" --strict
    python check_abbrev.py "论文.docx" --csv report.csv
"""
import re, sys, argparse, csv, json
from docx import Document


def find_abbrev_entries(docx_path):
    """
    扫描文档，找出所有疑似「中文（ABBR）」或「中文（English Name, ABBR）」的段落。
    返回列表，每项: {para_index, text, abbrev, full_name, has_full_name, first_occurrence}
    """
    doc = Document(docx_path)
    # 匹配 (ABBR) — 大写字母为主，可能有数字
    abbr_in_paren = re.compile(r'[（(]\s*([A-Z][A-Za-z0-9+/_-]{1,20})\s*[）)]')
    # 匹配带有英文全称的 (English Full Name, ABBR)
    full_with_abbr = re.compile(r'[（(]\s*([A-Za-z][A-Za-z\s.]+?)\s*[,，]\s*([A-Z][A-Za-z0-9+/_-]{1,20})\s*[）)]')

    results = []
    seen_abbrevs = {}  # abbrev -> first occurrence info
    known_ok = set()  # abbrevs already seen with full name

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Check for (Full Name, ABBR) pattern first
        for m in full_with_abbr.finditer(text):
            full_name = m.group(1).strip()
            abbr = m.group(2).strip()
            known_ok.add(abbr)
            if abbr not in seen_abbrevs:
                seen_abbrevs[abbr] = {
                    'para_index': i,
                    'text': text[:120],
                    'abbrev': abbr,
                    'full_name': full_name,
                    'has_full_name': True,
                    'first_occurrence': True,
                }

        # Check for bare (ABBR) pattern
        for m in abbr_in_paren.finditer(text):
            abbr = m.group(1).strip()
            if abbr in known_ok:
                continue
            # Check if this is actually a (Full, ABBR) that was already caught
            # by looking at the surrounding context
            start = max(0, m.start() - 5)
            context_before = text[start:m.start()]
            # If there's a comma before the abbreviation, it might be "Full, ABBR"
            if ',' in context_before or '，' in context_before:
                continue

            if abbr not in seen_abbrevs:
                # Check if full name is elsewhere in same paragraph before the paren
                before_text = text[:m.start()]
                has_full = bool(re.search(r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,5}', before_text))
                seen_abbrevs[abbr] = {
                    'para_index': i,
                    'text': text[:120],
                    'abbrev': abbr,
                    'full_name': '',
                    'has_full_name': has_full,
                    'first_occurrence': True,
                }

    # Mark non-first occurrences
    result_list = list(seen_abbrevs.values())
    return result_list


def check_abbrev_convention(docx_path, strict=False):
    """
    检查 docx 中所有英文简称的首次出现是否符合规范。
    规范：中文全称（English Full Name, ABBR）
    
    strict=True 时，即使英文全称出现在括号前也被认为不合规（要求括号内同时有全称和简称）。
    """
    entries = find_abbrev_entries(docx_path)
    
    issues = []
    ok = []
    for e in entries:
        if not e['has_full_name']:
            issues.append(e)
        elif strict:
            # In strict mode, the full name must be inside the parentheses with the abbreviation
            para = _get_para_text_by_index(docx_path, e['para_index'])
            pattern = re.compile(r'[（(]\s*' + re.escape(e['full_name']) + r'\s*[,，]\s*' + re.escape(e['abbrev']) + r'\s*[）)]')
            if not pattern.search(para):
                issues.append(e)
            else:
                ok.append(e)
        else:
            ok.append(e)
    
    return {'issues': issues, 'ok': ok}


def _get_para_text_by_index(docx_path, idx):
    from docx import Document
    doc = Document(docx_path)
    paras = doc.paragraphs
    if idx < len(paras):
        return paras[idx].text
    return ''


def main():
    parser = argparse.ArgumentParser(
        description='检查 docx 中英文简称首次出现是否符合中文（English Full Name, ABBR）规范')
    parser.add_argument('docx', help='输入的 docx 文件路径')
    parser.add_argument('-o', '--output', help='输出文件 (JSON 或 CSV，根据后缀判断)')
    parser.add_argument('--strict', action='store_true',
                        help='严格模式：要求括号内同时出现英文全称和简称')
    parser.add_argument('--csv', help='以 CSV 格式输出到指定文件')
    args = parser.parse_args()

    result = check_abbrev_convention(args.docx, strict=args.strict)

    issues = result['issues']
    ok = result['ok']

    print(f'检查完成。共发现 {len(ok)} 个合规简称，{len(issues)} 个不合规简称。\n')

    if issues:
        print('=== 不合规的英文简称（首次出现缺少英文全称）===\n')
        for e in issues:
            print(f'  段落 {e["para_index"]}: [{e["abbrev"]}]')
            print(f'  原文: {e["text"]}')
            print()
    else:
        print('所有英文简称首次出现均符合规范。\n')

    # 输出到文件
    if args.output:
        ext = args.output.rsplit('.', 1)[-1].lower()
        if ext == 'csv' or args.csv:
            _write_csv(args.output or args.csv, issues, ok)
        else:
            _write_json(args.output, issues, ok)
        print(f'详细报告已写入: {args.output}')


def _write_json(path, issues, ok):
    data = {'issues': issues, 'ok': ok,
            'summary': {'total_issues': len(issues), 'total_ok': len(ok)}}
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _write_csv(path, issues, ok):
    with open(path, 'w', encoding='utf-8-sig', newline='') as f:
        w = csv.writer(f)
        w.writerow(['status', 'para_index', 'abbrev', 'full_name', 'text'])
        for e in issues:
            w.writerow(['issue', e['para_index'], e['abbrev'], e.get('full_name', ''), e['text']])
        for e in ok:
            w.writerow(['ok', e['para_index'], e['abbrev'], e.get('full_name', ''), e['text']])


if __name__ == '__main__':
    main()
