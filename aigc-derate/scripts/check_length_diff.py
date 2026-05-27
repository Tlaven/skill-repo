"""
比较改写前后文档段落长度变化。检测字数偏差超过阈值的段落。

用法:
    python check_length_diff.py "改前备份.docx" "改后论文.docx"
    python check_length_diff.py "改前备份.docx" "改后论文.docx" --threshold 0.1
    python check_length_diff.py "改前备份.docx" "改后论文.docx" --json diff.json
"""
import json, sys, argparse
from docx import Document
from config import CONFIG


def check_length_diff(old_docx, new_docx, threshold=None):
    if threshold is None:
        threshold = CONFIG['length_diff_threshold']
    old_doc = Document(old_docx)
    new_doc = Document(new_docx)

    diffs = []
    max_len = min(len(old_doc.paragraphs), len(new_doc.paragraphs))

    for i in range(max_len):
        old_text = old_doc.paragraphs[i].text or ""
        new_text = new_doc.paragraphs[i].text or ""
        old_len = len(old_text)
        new_len = len(new_text)

        if old_len == 0 and new_len == 0:
            continue

        ratio = new_len / old_len if old_len > 0 else 1
        if abs(ratio - 1) > threshold:
            diffs.append({
                'paragraph': i,
                'old_chars': old_len,
                'new_chars': new_len,
                'ratio': round(ratio, 3),
                'old_preview': old_text[:60],
                'new_preview': new_text[:60]
            })

    return {
        'threshold': threshold,
        'total_checked': max_len,
        'total_over_threshold': len(diffs),
        'diffs': diffs
    }


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='检查改写前后段落字数偏差')
    parser.add_argument('old_docx', help='改写前备份的docx')
    parser.add_argument('new_docx', help='改写后的docx')
    parser.add_argument('--threshold', type=float, default=CONFIG['length_diff_threshold'],
                        help=f"偏差阈值（默认{CONFIG['length_diff_threshold']*100:.0f}%%）")
    parser.add_argument('--json', help='输出JSON结果到文件')
    args = parser.parse_args()

    result = check_length_diff(args.old_docx, args.new_docx, args.threshold)
    over = result['diffs']

    print(f"共检查 {result['total_checked']} 段")
    print(f"超过 {args.threshold:.0%} 偏差阈值的段落: {result['total_over_threshold']} 段\n")

    for d in sorted(over, key=lambda x: -abs(x['ratio'] - 1))[:20]:
        direction = "🔴缩短" if d['ratio'] < 1 else "🟡增长"
        print(f"P{d['paragraph']:3d} {direction}: {d['old_chars']}→{d['new_chars']}字 ({d['ratio']:+.0%})")
        print(f"  原文: {d['old_preview']}")
        print(f"  改后: {d['new_preview']}")
        print()

    if over:
        for d in over:
            if d['ratio'] < 0.5:
                print(f"⚠️ P{d['paragraph']} 严重缩短：{d['old_chars']}→{d['new_chars']}字")

    if args.json:
        with open(args.json, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
